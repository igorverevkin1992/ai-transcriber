"""Integration tests for API endpoints."""
import pytest
from fastapi.testclient import TestClient

from main import app


@pytest.fixture
def client():
    with TestClient(app) as c:
        yield c


class TestHealthEndpoint:
    def test_health_returns_ok(self, client):
        resp = client.get("/health")
        assert resp.status_code == 200
        data = resp.json()
        assert data["status"] == "ok"
        assert data["service"] == "ABTGS Backend"


class TestCreateProject:
    def test_valid_url(self, client):
        resp = client.post("/api/v1/projects", json={"url": "https://yadi.sk/d/test123"})
        assert resp.status_code == 200
        data = resp.json()
        assert "id" in data
        assert len(data["id"]) == 36  # UUID

    def test_invalid_url_host(self, client):
        resp = client.post("/api/v1/projects", json={"url": "https://evil.com/file"})
        assert resp.status_code == 400
        assert "Яндекс.Диск" in resp.json()["detail"]

    def test_missing_url(self, client):
        resp = client.post("/api/v1/projects", json={})
        assert resp.status_code == 422


class TestUploadWithPassport:
    def _docx_bytes(self):
        import io

        from docx import Document
        doc = Document()
        doc.add_paragraph("Герои: Иванов Иван")
        doc.add_paragraph("Количество героев: 1")
        buf = io.BytesIO()
        doc.save(buf)
        return buf.getvalue()

    def test_passport_threaded_into_task(self, client, monkeypatch):
        from pathlib import Path

        from backend import routes, services

        captured = {}
        monkeypatch.setattr(routes, "validate_mime_type", lambda b: None)
        monkeypatch.setattr(routes, "submit_task", lambda *a, **k: captured.update(k))

        resp = client.post(
            "/api/v1/batch/upload",
            data={"engine": "whisper", "whisper_model": "medium"},
            files={
                "file": ("video.mp4", b"\x00" * 64, "video/mp4"),
                "passport": ("passport.docx", self._docx_bytes(),
                             "application/vnd.openxmlformats-officedocument.wordprocessingml.document"),
            },
        )
        assert resp.status_code == 200
        pid = resp.json()["id"]
        assert captured.get("passport_path", "").endswith("_passport.docx")
        passport_file = Path(captured["passport_path"])
        assert passport_file.exists()
        assert services.projects_db.get(pid)["task_kwargs"]["passport_path"] == str(passport_file)
        # cleanup
        passport_file.unlink(missing_ok=True)
        Path(services.TEMP_DIR / f"{pid}_video.mp4").unlink(missing_ok=True)

    def test_passport_wrong_extension_rejected(self, client, monkeypatch):
        from backend import routes

        monkeypatch.setattr(routes, "validate_mime_type", lambda b: None)
        monkeypatch.setattr(routes, "submit_task", lambda *a, **k: None)
        resp = client.post(
            "/api/v1/batch/upload",
            data={"engine": "whisper", "whisper_model": "medium"},
            files={
                "file": ("video.mp4", b"\x00" * 64, "video/mp4"),
                "passport": ("passport.txt", b"hello", "text/plain"),
            },
        )
        assert resp.status_code == 400


class TestProjectStatus:
    def test_not_found(self, client):
        resp = client.get("/api/v1/projects/nonexistent/status")
        assert resp.status_code == 404

    def test_status_after_create(self, client):
        create_resp = client.post("/api/v1/projects", json={"url": "https://yadi.sk/d/test"})
        pid = create_resp.json()["id"]
        status_resp = client.get(f"/api/v1/projects/{pid}/status")
        assert status_resp.status_code == 200
        data = status_resp.json()
        assert data["status"] in ("queued", "downloading", "error")
        assert "status_label" in data


class TestProjectResult:
    def test_not_found(self, client):
        resp = client.get("/api/v1/projects/nonexistent")
        assert resp.status_code == 404

    def test_not_completed(self, client):
        create_resp = client.post("/api/v1/projects", json={"url": "https://yadi.sk/d/test"})
        pid = create_resp.json()["id"]
        resp = client.get(f"/api/v1/projects/{pid}")
        assert resp.status_code == 400


class TestExport:
    def test_not_found(self, client):
        resp = client.post(
            "/api/v1/projects/nonexistent/export",
            json={"mappings": [], "filename": "test.docx"},
        )
        assert resp.status_code == 404
