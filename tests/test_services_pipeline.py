"""Tests for core pipeline functions in backend.services.

These cover the deterministic units (result processing, cancellation,
auto-export, retry scheduling, executor lifecycle) with the heavy
Whisper/SpeechKit/ffmpeg dependencies mocked out or avoided.
"""

from pathlib import Path

import pytest

from backend import services
from backend.models import ProjectStatusEnum
from backend.store import ProjectStore


@pytest.fixture
def fresh_store(tmp_path, monkeypatch):
    """Swap services.projects_db with an isolated store + temp/output dirs."""
    store = ProjectStore(db_path=str(tmp_path / "svc.db"))
    temp_dir = tmp_path / "temp"
    out_dir = tmp_path / "out"
    temp_dir.mkdir()
    out_dir.mkdir()
    monkeypatch.setattr(services, "projects_db", store)
    monkeypatch.setattr(services, "TEMP_DIR", temp_dir)
    monkeypatch.setattr(services, "OUTPUT_DIR", out_dir)
    return store, temp_dir, out_dir


def _seg(channel, text, start_ms, end_ms):
    return {
        "channel_tag": channel,
        "text": text,
        "words": [{"text": text, "start_ms": start_ms, "end_ms": end_ms}],
    }


class TestProcessRecognitionResult:
    def test_builds_segments_speakers_and_timecodes(self, fresh_store):
        store, _, _ = fresh_store
        store.create("p1", {"id": "p1", "status": ProjectStatusEnum.TRANSCRIBING, "created_at": 1.0})
        segments = [
            _seg("0", "Привет.", 0, 1000),
            _seg("1", "Здравствуйте.", 1000, 3000),
            _seg("0", "Как дела?", 3000, 4000),
        ]
        services._process_recognition_result("p1", segments, "Иванов_Петров.mp4", Path("/nonexistent.mp4"))

        result = store["p1"]["result"]
        assert len(result["segments"]) == 3
        assert result["segments"][0]["timecode"] == "00:00:00:00"
        # speaker "1" speaks 2s vs "0" 2s total — both present
        assert set(result["speakers"].keys()) == {"0", "1"}
        assert store["p1"]["fps"] == 25  # fallback when video missing

    def test_filters_empty_text_segments(self, fresh_store):
        store, _, _ = fresh_store
        store.create("p2", {"id": "p2", "status": ProjectStatusEnum.TRANSCRIBING, "created_at": 1.0})
        segments = [
            _seg("0", "Реальный текст.", 0, 1000),
            _seg("0", "   ", 1000, 2000),  # whitespace-only -> dropped
            {"channel_tag": "0", "text": "Без слов", "words": []},  # no words -> dropped
        ]
        services._process_recognition_result("p2", segments, "file.mp4", Path("/nope.mp4"))
        assert len(store["p2"]["result"]["segments"]) == 1

    def test_speaker_names_from_filename(self, fresh_store):
        store, _, _ = fresh_store
        store.create("p3", {"id": "p3", "status": ProjectStatusEnum.TRANSCRIBING, "created_at": 1.0})
        # speaker 0 talks longest -> gets first filename name
        segments = [
            _seg("0", "много текста тут.", 0, 10000),
            _seg("1", "мало.", 10000, 11000),
        ]
        services._process_recognition_result("p3", segments, "Маданов_Антипенко.mp4", Path("/nope.mp4"))
        speakers = store["p3"]["result"]["speakers"]
        assert speakers["0"]["suggested_name"] == "Маданов"
        assert speakers["1"]["suggested_name"] == "Антипенко"

    def test_single_speaker_sets_low_confidence(self, fresh_store):
        store, _, _ = fresh_store
        store.create("p4", {"id": "p4", "status": ProjectStatusEnum.TRANSCRIBING, "created_at": 1.0})
        services._process_recognition_result(
            "p4", [_seg("0", "Один голос.", 0, 1000)], "file.mp4", Path("/nope.mp4")
        )
        assert store["p4"]["result"]["low_confidence_diarization"] is True

    def test_single_speaker_adds_warning(self, fresh_store):
        store, _, _ = fresh_store
        store.create("pw1", {"id": "pw1", "status": ProjectStatusEnum.TRANSCRIBING, "created_at": 1.0})
        services._process_recognition_result(
            "pw1", [_seg("0", "Один голос.", 0, 1000)], "file.mp4", Path("/nope.mp4")
        )
        warnings = store["pw1"]["result"]["warnings"]
        assert any("1 говорящий" in w for w in warnings)

    def test_missing_start_tc_adds_warning(self, fresh_store):
        store, _, _ = fresh_store
        store.create("pw2", {"id": "pw2", "status": ProjectStatusEnum.TRANSCRIBING, "created_at": 1.0})
        # filename without TC -> start 00:00:00:00 -> warning
        services._process_recognition_result(
            "pw2", [_seg("0", "Текст.", 0, 1000), _seg("1", "Ответ.", 1000, 2000)],
            "file.mp4", Path("/nope.mp4"),
        )
        warnings = store["pw2"]["result"]["warnings"]
        assert any("таймкод" in w.lower() for w in warnings)

    def test_start_tc_from_filename_no_warning(self, fresh_store):
        store, _, _ = fresh_store
        store.create("pw3", {"id": "pw3", "status": ProjectStatusEnum.TRANSCRIBING, "created_at": 1.0})
        services._process_recognition_result(
            "pw3", [_seg("0", "Текст.", 0, 1000), _seg("1", "Ответ.", 1000, 2000)],
            "Иванов_Петров_04.41.18.00_f21.mp4", Path("/nope.mp4"),
        )
        result = store["pw3"]["result"]
        assert result["segments"][0]["timecode"].startswith("04:41:18")
        assert not any("таймкод" in w.lower() for w in result["warnings"])


class TestBuildWhisperPrompt:
    def test_base_prompt_only(self, monkeypatch):
        monkeypatch.setattr(services, "TRANSCRIPT_GLOSSARY", "")
        # date + file-code only -> no speaker names extracted
        prompt = services._build_whisper_prompt("09.06.2026_f21.mp4")
        assert prompt == services.WHISPER_INITIAL_PROMPT

    def test_names_from_filename(self, monkeypatch):
        monkeypatch.setattr(services, "TRANSCRIPT_GLOSSARY", "")
        prompt = services._build_whisper_prompt("Иванов_Петров.mp4")
        assert "Участники: Иванов, Петров." in prompt

    def test_glossary_injected(self, monkeypatch):
        monkeypatch.setattr(services, "TRANSCRIPT_GLOSSARY", "Мордюкова, star quality")
        prompt = services._build_whisper_prompt("09.06.2026_f21.mp4")
        assert "Мордюкова, star quality" in prompt

    def test_autodetect_interviewer_and_name_mapping(self, fresh_store):
        store, _, _ = fresh_store
        store.create("azk1", {"id": "azk1", "status": ProjectStatusEnum.TRANSCRIBING, "created_at": 1.0})
        # speaker "0" = interviewer alternating with guests 1, 2, 3 (3 interviews)
        segs = [
            _seg("0", "Вопрос раз.", 0, 1000),
            _seg("1", "Ответ.", 1000, 3000),
            _seg("0", "Вопрос два.", 3000, 4000),
            _seg("1", "Ещё ответ.", 4000, 6000),
            _seg("0", "Вопрос три.", 6000, 7000),
            _seg("2", "Ответ гостя два.", 7000, 9000),
            _seg("0", "Вопрос четыре.", 9000, 10000),
            _seg("2", "Снова ответ.", 10000, 12000),
            _seg("0", "Вопрос пять.", 12000, 13000),
            _seg("3", "Ответ гостя три.", 13000, 15000),
            _seg("0", "Вопрос шесть.", 15000, 16000),
            _seg("3", "Финал.", 16000, 18000),
        ]
        services._process_recognition_result(
            "azk1", segs, "Кравченко_Артемьева_Нилов_f21.mp4", Path("/nope.mp4")
        )
        speakers = store["azk1"]["result"]["speakers"]
        assert speakers["0"]["suggested_name"] == "АЗК"
        assert speakers["1"]["suggested_name"] == "Кравченко"
        assert speakers["2"]["suggested_name"] == "Артемьева"
        assert speakers["3"]["suggested_name"] == "Нилов"

    def test_extra_warnings_propagate_to_result(self, fresh_store):
        store, _, _ = fresh_store
        store.create("w1", {"id": "w1", "status": ProjectStatusEnum.TRANSCRIBING, "created_at": 1.0})
        segs = [_seg("0", "Привет.", 0, 1000), _seg("1", "Здравствуйте.", 1500, 2500)]
        services._process_recognition_result(
            "w1", segs, "file.mp4", Path("/nope.mp4"), extra_warnings=["тестовое предупреждение"]
        )
        assert "тестовое предупреждение" in store["w1"]["result"]["warnings"]

    def test_merges_same_speaker_turns(self, fresh_store):
        store, _, _ = fresh_store
        store.create("pm1", {"id": "pm1", "status": ProjectStatusEnum.TRANSCRIBING, "created_at": 1.0})
        segments = [
            _seg("0", "Привет.", 0, 1000),
            _seg("0", "Как дела?", 1500, 2500),
            _seg("1", "Хорошо.", 3000, 4000),
        ]
        services._process_recognition_result("pm1", segments, "file.mp4", Path("/nope.mp4"))
        result = store["pm1"]["result"]["segments"]
        assert len(result) == 2
        assert result[0]["text"] == "Привет. Как дела?"
        assert result[1]["text"] == "Хорошо."

    def test_speaker_durations_unaffected_by_merge(self, fresh_store):
        store, _, _ = fresh_store
        store.create("pm2", {"id": "pm2", "status": ProjectStatusEnum.TRANSCRIBING, "created_at": 1.0})
        segments = [
            _seg("0", "Привет.", 0, 1000),
            _seg("0", "Как дела?", 1500, 2500),
            _seg("1", "Хорошо.", 3000, 4000),
        ]
        services._process_recognition_result("pm2", segments, "file.mp4", Path("/nope.mp4"))
        speakers = store["pm2"]["result"]["speakers"]
        assert speakers["0"]["duration_sec"] == 2.0
        assert speakers["1"]["duration_sec"] == 1.0

    def test_turn_merge_disabled(self, fresh_store, monkeypatch):
        store, _, _ = fresh_store
        monkeypatch.setattr(services, "TURN_MERGE_ENABLED", False)
        store.create("pm3", {"id": "pm3", "status": ProjectStatusEnum.TRANSCRIBING, "created_at": 1.0})
        segments = [
            _seg("0", "Привет.", 0, 1000),
            _seg("0", "Как дела?", 1500, 2500),
        ]
        services._process_recognition_result("pm3", segments, "file.mp4", Path("/nope.mp4"))
        assert len(store["pm3"]["result"]["segments"]) == 2

    def test_tech_break_inserted_on_long_gap(self, fresh_store):
        store, _, _ = fresh_store
        store.create("pm4", {"id": "pm4", "status": ProjectStatusEnum.TRANSCRIBING, "created_at": 1.0})
        segments = [
            _seg("0", "До перерыва.", 0, 5000),
            _seg("0", "После перерыва.", 50000, 55000),
        ]
        services._process_recognition_result("pm4", segments, "file.mp4", Path("/nope.mp4"))
        result = store["pm4"]["result"]["segments"]
        assert [s["text"] for s in result] == [
            "До перерыва.", "(Технические моменты).", "После перерыва.",
        ]
        assert result[1]["timecode"] == "00:00:05:00"


class TestMakeDiarizePipeline:
    """Конструктор DiarizationPipeline переименовал use_auth_token -> token
    между версиями whisperx; помощник должен подобрать правильный аргумент и
    закрепить модель диаризации (DIARIZATION_MODEL)."""

    def test_uses_token_kwarg(self, monkeypatch):
        monkeypatch.setattr(services, "DIARIZATION_MODEL", "pyannote/speaker-diarization-3.1")
        captured = {}

        class FakePipeline:
            def __init__(self, model_name=None, token=None, device="cpu"):
                captured["model_name"] = model_name
                captured["token"] = token
                captured["device"] = device

        services._make_diarize_pipeline(FakePipeline, "hf_abc", "cpu")
        assert captured == {
            "model_name": "pyannote/speaker-diarization-3.1",
            "token": "hf_abc",
            "device": "cpu",
        }

    def test_uses_use_auth_token_kwarg(self, monkeypatch):
        monkeypatch.setattr(services, "DIARIZATION_MODEL", "pyannote/speaker-diarization-3.1")
        captured = {}

        class FakePipeline:
            def __init__(self, model_name=None, use_auth_token=None, device="cpu"):
                captured["model_name"] = model_name
                captured["use_auth_token"] = use_auth_token
                captured["device"] = device

        services._make_diarize_pipeline(FakePipeline, "hf_xyz", "cpu")
        assert captured == {
            "model_name": "pyannote/speaker-diarization-3.1",
            "use_auth_token": "hf_xyz",
            "device": "cpu",
        }

    def test_no_model_name_param(self, monkeypatch):
        """Старый конструктор без model_name — модель не навязываем."""
        monkeypatch.setattr(services, "DIARIZATION_MODEL", "pyannote/speaker-diarization-3.1")
        captured = {}

        class FakePipeline:
            def __init__(self, use_auth_token=None, device="cpu"):
                captured["use_auth_token"] = use_auth_token
                captured["device"] = device

        services._make_diarize_pipeline(FakePipeline, "hf_old", "cpu")
        assert captured == {"use_auth_token": "hf_old", "device": "cpu"}


class TestIsCudaError:
    """Детектор CUDA-сбоев для отката транскрипции на CPU."""

    def test_invalid_device_ordinal(self):
        exc = RuntimeError("parallel_for failed: cudaErrorInvalidDevice: invalid device ordinal")
        assert services._is_cuda_error(exc)

    def test_cudnn_error(self):
        assert services._is_cuda_error(RuntimeError("Could not load library cudnn_ops64_9.dll"))

    def test_unrelated_error_not_matched(self):
        assert not services._is_cuda_error(RuntimeError("Файл не найден"))


class TestProcessVideoTask:
    def test_runs_postprocessing(self, fresh_store, monkeypatch):
        """Регрессия: SpeechKit-путь (Яндекс.Диск) должен прогонять постобработку."""
        store, _, _ = fresh_store
        store.create("pv1", {"id": "pv1", "status": ProjectStatusEnum.QUEUED, "created_at": 1.0})

        calls = {}
        monkeypatch.setattr(services, "_download_from_yadisk",
                            lambda pid, url, path: "Иванов_Петров.mp4")
        monkeypatch.setattr(services, "_convert_to_opus", lambda pid, i, o: None)
        monkeypatch.setattr(services, "_transcribe_with_speechkit",
                            lambda pid, path: [_seg("0", "привет.", 0, 1000)])

        import backend.postprocess as pp

        def fake_postprocess(segments, use_gemini=True, warnings=None):
            calls["postprocess"] = True
            return segments

        monkeypatch.setattr(pp, "postprocess_segments", fake_postprocess)

        services.process_video_task("pv1", "http://disk")

        assert calls.get("postprocess") is True
        assert store["pv1"]["status"] == ProjectStatusEnum.COMPLETED


class TestCancelProject:
    def test_returns_false_for_missing(self, fresh_store):
        assert services.cancel_project("ghost") is False

    def test_removes_project_and_temp_files(self, fresh_store):
        store, temp_dir, _ = fresh_store
        store.create("p5", {"id": "p5", "status": ProjectStatusEnum.QUEUED, "created_at": 1.0})
        # create matching temp artifacts
        (temp_dir / "p5_video.mp4").write_bytes(b"x")
        (temp_dir / "p5.opus").write_bytes(b"y")

        assert services.cancel_project("p5") is True
        assert "p5" not in store
        assert not (temp_dir / "p5_video.mp4").exists()
        assert not (temp_dir / "p5.opus").exists()


class TestAutoExportProject:
    def test_returns_none_without_result(self, fresh_store):
        store, _, _ = fresh_store
        store.create("p6", {"id": "p6", "status": ProjectStatusEnum.QUEUED, "created_at": 1.0})
        assert services.auto_export_project("p6", "/tmp/x.docx") is None

    def test_generates_docx_and_excludes_tech_speakers(self, fresh_store):
        store, temp_dir, _ = fresh_store
        store.create("p7", {
            "id": "p7",
            "status": ProjectStatusEnum.COMPLETED,
            "created_at": 1.0,
            "original_filename": "interview.mp4",
            "result": {
                "speakers": {
                    "0": {"duration_sec": 100.0, "suggested_name": "Иванов"},
                    "1": {"duration_sec": 50.0, "suggested_name": "АЗК"},
                },
                "segments": [
                    {"timecode": "00:00:01:00", "speaker": "0", "text": "Текст."},
                    {"timecode": "00:00:05:00", "speaker": "1", "text": "Вопрос."},
                ],
            },
        })
        out = str(temp_dir / "p7.docx")
        name = services.auto_export_project("p7", out)
        assert name == "interview.docx"
        assert Path(out).exists()

        from docx import Document
        doc = Document(out)
        legend = [p.text for p in doc.paragraphs if "–" in p.text]
        assert not any("АЗК" in t for t in legend)  # tech speaker excluded from legend
        assert any("Иванов" in t for t in legend)


    def test_service_labels_keep_full_prefix_and_dont_pollute_abbrs(self, fresh_store):
        store, temp_dir, _ = fresh_store
        store.create("p7b", {
            "id": "p7b",
            "status": ProjectStatusEnum.COMPLETED,
            "created_at": 1.0,
            "original_filename": "interview.mp4",
            "result": {
                "speakers": {
                    "0": {"duration_sec": 100.0, "suggested_name": "Антипенко"},
                    "1": {"duration_sec": 50.0, "suggested_name": "АЗК"},
                    "2": {"duration_sec": 30.0, "suggested_name": "Интервьюер"},
                },
                "segments": [
                    {"timecode": "00:00:01:00", "speaker": "0", "text": "Текст."},
                    {"timecode": "00:00:05:00", "speaker": "1", "text": "Вопрос."},
                    {"timecode": "00:00:10:00", "speaker": "2", "text": "Ответ."},
                ],
            },
        })
        out = str(temp_dir / "p7b.docx")
        services.auto_export_project("p7b", out)

        from docx import Document
        doc = Document(out)
        all_text = "\n".join(p.text for p in doc.paragraphs)
        # АЗК uses its full name as prefix, not a computed abbreviation
        assert "АЗК:" in all_text
        # Антипенко gets clean "А" (no collision with АЗК)
        assert "А:" in all_text


    def test_legend_name_inverted_to_first_name_first(self, fresh_store):
        store, temp_dir, _ = fresh_store
        store.create("p7c", {
            "id": "p7c",
            "status": ProjectStatusEnum.COMPLETED,
            "created_at": 1.0,
            "original_filename": "Довлатова Алла_interview.mp4",
            "result": {
                "speakers": {
                    "0": {"duration_sec": 100.0, "suggested_name": "Довлатова Алла"},
                    "1": {"duration_sec": 50.0, "suggested_name": "Интервьюер"},
                },
                "segments": [
                    {"timecode": "00:00:01:00", "speaker": "0", "text": "Текст."},
                    {"timecode": "00:00:05:00", "speaker": "1", "text": "Вопрос."},
                ],
            },
        })
        out = str(temp_dir / "p7c.docx")
        services.auto_export_project("p7c", out)

        from docx import Document
        doc = Document(out)
        legend = [p.text for p in doc.paragraphs if "–" in p.text]
        assert any("Алла Довлатова" in t for t in legend)
        assert not any("Довлатова Алла" in t for t in legend)
        all_text = "\n".join(p.text for p in doc.paragraphs)
        assert "Д:" in all_text


class TestMaybeRetry:
    def test_no_retry_when_not_error(self, fresh_store):
        store, _, _ = fresh_store
        store.create("p8", {"id": "p8", "status": ProjectStatusEnum.COMPLETED, "created_at": 1.0})
        services._maybe_retry("p8")
        assert store["p8"]["status"] == ProjectStatusEnum.COMPLETED

    def test_requeues_failed_project(self, fresh_store, monkeypatch):
        store, _, _ = fresh_store
        # capture Timer instead of really scheduling
        scheduled = {}

        class FakeTimer:
            def __init__(self, delay, func, args=(), kwargs=None):
                scheduled["delay"] = delay
                scheduled["args"] = args
            def start(self):
                scheduled["started"] = True

        monkeypatch.setattr(services.threading, "Timer", FakeTimer)
        store.create("p9", {
            "id": "p9",
            "status": ProjectStatusEnum.ERROR,
            "created_at": 1.0,
            "retry_count": 0,
            "task_func": "process_video_task",
            "task_args": ("p9", "http://disk"),
        })
        services._maybe_retry("p9")
        assert store["p9"]["status"] == ProjectStatusEnum.QUEUED
        assert store["p9"]["retry_count"] == 1
        assert scheduled.get("started") is True

    def test_stops_at_max_retries(self, fresh_store, monkeypatch):
        store, _, _ = fresh_store
        monkeypatch.setattr(services, "MAX_RETRIES", 3)
        store.create("p10", {
            "id": "p10",
            "status": ProjectStatusEnum.ERROR,
            "created_at": 1.0,
            "retry_count": 3,
            "task_func": "process_video_task",
            "task_args": ("p10", "http://disk"),
        })
        services._maybe_retry("p10")
        # stays ERROR, no requeue
        assert store["p10"]["status"] == ProjectStatusEnum.ERROR
        assert store["p10"]["retry_count"] == 3


class _SyncExecutor:
    """Исполняет submit(fn) синхронно — для детерминированных тестов submit_task."""
    def submit(self, fn):
        fn()
        class _F:
            def result(self, timeout=None):
                return None
        return _F()


class TestSubmitTaskIdempotency:
    def test_skips_completed_project(self, fresh_store, monkeypatch):
        # Дубликат-отправка завершённого проекта не должна запускать обработку
        # повторно (иначе падает «Файл не найден» и затирает COMPLETED).
        store, _, _ = fresh_store
        monkeypatch.setattr(services, "_ensure_executor", lambda: _SyncExecutor())
        store.create("c1", {"id": "c1", "status": ProjectStatusEnum.COMPLETED, "created_at": 1.0})
        calls = []
        services.submit_task(lambda *a, **k: calls.append(1), "c1", project_id="c1")
        assert calls == []
        assert store["c1"]["status"] == ProjectStatusEnum.COMPLETED

    def test_runs_queued_project(self, fresh_store, monkeypatch):
        store, _, _ = fresh_store
        monkeypatch.setattr(services, "_ensure_executor", lambda: _SyncExecutor())
        store.create("c2", {"id": "c2", "status": ProjectStatusEnum.QUEUED, "created_at": 1.0})
        calls = []
        services.submit_task(lambda *a, **k: calls.append(1), "c2", project_id="c2")
        assert calls == [1]
        # после завершения проект убран из набора «выполняющихся»
        assert "c2" not in services._running_projects

    def test_skips_when_already_running(self, fresh_store, monkeypatch):
        store, _, _ = fresh_store
        monkeypatch.setattr(services, "_ensure_executor", lambda: _SyncExecutor())
        store.create("c3", {"id": "c3", "status": ProjectStatusEnum.QUEUED, "created_at": 1.0})
        services._running_projects.add("c3")
        try:
            calls = []
            services.submit_task(lambda *a, **k: calls.append(1), "c3", project_id="c3")
            assert calls == []
        finally:
            services._running_projects.discard("c3")

    def test_resume_refuses_completed(self, fresh_store):
        store, _, _ = fresh_store
        store.create("c4", {
            "id": "c4", "status": ProjectStatusEnum.COMPLETED, "created_at": 1.0,
            "task_func": "process_video_task", "task_args": ("c4", "http://disk"),
        })
        ok, msg = services.resume_project("c4")
        assert ok is False
        assert "завершён" in msg

    def test_resume_refuses_running(self, fresh_store):
        store, _, _ = fresh_store
        store.create("c5", {
            "id": "c5", "status": ProjectStatusEnum.ERROR, "created_at": 1.0,
            "task_func": "process_video_task", "task_args": ("c5", "http://disk"),
        })
        services._running_projects.add("c5")
        try:
            ok, msg = services.resume_project("c5")
            assert ok is False
            assert "выполняется" in msg
        finally:
            services._running_projects.discard("c5")


class TestEnsureExecutor:
    def test_recreates_after_shutdown(self, monkeypatch):
        # simulate a shut-down executor
        services.shutdown_executor()
        assert services._executor_alive is False
        ex = services._ensure_executor()
        assert services._executor_alive is True
        # submit a trivial task to prove it's live
        fut = ex.submit(lambda: 42)
        assert fut.result(timeout=5) == 42


class TestEventLevelBoundaryCorrection:
    def test_glued_turn_split_after_correction(self, fresh_store, monkeypatch):
        # Событие «Вы готовы?» ошибочно помечено гостем: правка на уровне
        # событий переназначает его ведущему, и пересборка реплик разрезает
        # склейку (реплика гостя больше не содержит чужой вопрос).
        store, _, _ = fresh_store
        store.create("ev1", {"id": "ev1", "status": ProjectStatusEnum.TRANSCRIBING, "created_at": 1.0})
        monkeypatch.setattr(services, "SPEAKER_BOUNDARY_CORRECTION", True)
        monkeypatch.setattr(services, "INTERVIEWER_AUTODETECT", False)

        import backend.postprocess as pp
        monkeypatch.setattr(pp, "_gemini_call", lambda prompt, **kw: '[{"id": 1, "speaker": "0"}]')

        segments = [
            {"channel_tag": "1", "text": "Длинный ответ гостя.",
             "words": [{"text": "Длинный ответ гостя.", "start_ms": 0, "end_ms": 2000}]},
            {"channel_tag": "1", "text": "Вы готовы?",
             "words": [{"text": "Вы готовы?", "start_ms": 2500, "end_ms": 3500}]},
            {"channel_tag": "0", "text": "Да, готов.",
             "words": [{"text": "Да, готов.", "start_ms": 4000, "end_ms": 5000}]},
        ]
        services._process_recognition_result("ev1", segments, "file.mp4", Path("/nope.mp4"))

        result = store["ev1"]["result"]
        texts = [(s["speaker"], s["text"]) for s in result["segments"]]
        # Вопрос ушёл из реплики гостя и склеился с ответом спикера 0.
        assert ("1", "Длинный ответ гостя.") in texts
        assert any(sp == "0" and "Вы готовы?" in t and "Да, готов." in t for sp, t in texts)


class TestRecognitionCheckpoint:
    SEGS = [
        {"channel_tag": "0", "text": "Привет.",
         "words": [{"text": "Привет.", "start_ms": 0, "end_ms": 1000}]},
        {"channel_tag": "1", "text": "Здравствуйте.",
         "words": [{"text": "Здравствуйте.", "start_ms": 1500, "end_ms": 2500}]},
    ]

    def test_roundtrip(self, fresh_store):
        services._save_recognition_checkpoint("cp1", "whisper", "large", "f.wmv", self.SEGS)
        assert services._load_recognition_checkpoint("cp1", "whisper", "large", "f.wmv") == self.SEGS

    def test_mismatch_ignored(self, fresh_store):
        services._save_recognition_checkpoint("cp2", "whisper", "large", "f.wmv", self.SEGS)
        # Другая модель или другой файл — чекпоинт не подходит.
        assert services._load_recognition_checkpoint("cp2", "whisper", "medium", "f.wmv") is None
        assert services._load_recognition_checkpoint("cp2", "whisper", "large", "other.wmv") is None

    def test_corrupted_ignored(self, fresh_store):
        services._recognition_checkpoint_path("cp3").write_text("{оборвано", encoding="utf-8")
        assert services._load_recognition_checkpoint("cp3", "whisper", "large", "f.wmv") is None

    def test_retry_skips_transcription(self, fresh_store, monkeypatch):
        """Падение в постобработке → повтор берёт чекпоинт и НЕ гоняет Whisper заново."""
        import time as _time
        store, temp_dir, _ = fresh_store
        # created_at свежий: между двумя запусками проект в статусе ERROR не
        # должен попасть под TTL-очистку в начале второго запуска.
        store.create("cp4", {"id": "cp4", "status": ProjectStatusEnum.QUEUED,
                             "created_at": _time.time()})
        video = temp_dir / "cp4_video.wmv"
        video.write_bytes(b"x")

        transcribe_calls = []

        def fake_transcribe(pid, path, model, **kw):
            transcribe_calls.append(model)
            return list(self.SEGS)

        monkeypatch.setattr(services, "WHISPERX_AVAILABLE", True)
        monkeypatch.setattr(services, "_transcribe_with_whisperx", fake_transcribe)
        monkeypatch.setattr(services, "OCR_TIMECODE", False)

        import backend.postprocess as pp
        state = {"fail": True}

        def fake_postprocess(segments, warnings=None, crew_names=None, description=None):
            if state["fail"]:
                raise RuntimeError("Gemini упал")
            return segments

        monkeypatch.setattr(pp, "postprocess_segments", fake_postprocess)

        services.process_uploaded_file_task("cp4", video, "f.wmv",
                                            engine="whisper", whisper_model="large")
        assert store["cp4"]["status"] == ProjectStatusEnum.ERROR
        # Чекпоинт и исходник сохранены для повтора.
        assert services._recognition_checkpoint_path("cp4").exists()
        assert video.exists()

        state["fail"] = False
        services.process_uploaded_file_task("cp4", video, "f.wmv",
                                            engine="whisper", whisper_model="large")
        assert store["cp4"]["status"] == ProjectStatusEnum.COMPLETED
        assert transcribe_calls == ["large"]  # Whisper вызван ровно один раз
        assert not services._recognition_checkpoint_path("cp4").exists()

    def test_final_failure_drops_checkpoint(self, fresh_store, monkeypatch):
        """Исчерпаны ретраи → чекпоинт не должен переживать задачу."""
        store, temp_dir, _ = fresh_store
        store.create("cp5", {"id": "cp5", "status": ProjectStatusEnum.QUEUED,
                             "created_at": 1.0, "retry_count": services.MAX_RETRIES})
        video = temp_dir / "cp5_video.wmv"
        video.write_bytes(b"x")
        services._save_recognition_checkpoint("cp5", "whisper", "large", "f.wmv",
                                              self.SEGS)

        import backend.postprocess as pp
        monkeypatch.setattr(pp, "postprocess_segments",
                            lambda *a, **k: (_ for _ in ()).throw(RuntimeError("упал")))

        services.process_uploaded_file_task("cp5", video, "f.wmv",
                                            engine="whisper", whisper_model="large")
        assert store["cp5"]["status"] == ProjectStatusEnum.ERROR
        assert not services._recognition_checkpoint_path("cp5").exists()
        assert not video.exists()


class TestStageTimerAndDetail:
    def test_timer_logs_minutes_and_ratio(self, caplog):
        import logging
        with caplog.at_level(logging.INFO):
            with services._stage_timer("proj4242", "тестовый", audio_seconds=60):
                pass
        msgs = [r.getMessage() for r in caplog.records]
        assert any("Этап «тестовый»" in m and "длительности аудио" in m for m in msgs)

    def test_timer_logs_on_exception(self, caplog):
        import logging
        with caplog.at_level(logging.INFO):
            with pytest.raises(ValueError):
                with services._stage_timer("proj4242", "упавший"):
                    raise ValueError("boom")
        assert any("Этап «упавший»" in r.getMessage() for r in caplog.records)

    def test_set_stage_updates_fields(self, fresh_store):
        store, _, _ = fresh_store
        store.create("sd1", {"id": "sd1", "status": ProjectStatusEnum.TRANSCRIBING, "created_at": 1.0})
        services._set_stage("sd1", "Определение говорящих (диаризация)…", 65)
        assert store["sd1"]["status_detail"] == "Определение говорящих (диаризация)…"
        assert store["sd1"]["progress_percent"] == 65

    def test_set_stage_missing_project_noop(self, fresh_store):
        services._set_stage("ghost", "Этап…", 10)  # не должно бросать

    def test_retry_sets_visible_detail(self, fresh_store, monkeypatch):
        store, _, _ = fresh_store

        class FakeTimer:
            def __init__(self, delay, func, args=(), kwargs=None):
                pass

            def start(self):
                pass

        monkeypatch.setattr(services.threading, "Timer", FakeTimer)
        store.create("sd2", {
            "id": "sd2", "status": ProjectStatusEnum.ERROR, "created_at": 1.0,
            "retry_count": 0, "task_func": "process_video_task",
            "task_args": ("sd2", "http://disk"),
        })
        services._maybe_retry("sd2")
        assert "Повтор 1/" in store["sd2"]["status_detail"]


class TestDocTitleFlag:
    def _project(self, store):
        store.create("t1", {
            "id": "t1", "status": ProjectStatusEnum.COMPLETED, "created_at": 1.0,
            "original_filename": "интервью_ф13.wmv",
            "result": {
                "speakers": {"0": {"duration_sec": 10.0, "suggested_name": "Иванов"}},
                "segments": [{"timecode": "00:00:01:00", "speaker": "0", "text": "Текст."}],
            },
        })

    def test_title_present_by_default(self, fresh_store):
        store, temp_dir, _ = fresh_store
        self._project(store)
        out = str(temp_dir / "t1.docx")
        services.auto_export_project("t1", out)
        from docx import Document
        texts = [p.text for p in Document(out).paragraphs]
        assert any("интервью_ф13" in t for t in texts)

    def test_title_hidden_when_disabled(self, fresh_store, monkeypatch):
        import backend.docx_export as dx
        monkeypatch.setattr(dx, "DOC_TITLE_ENABLED", False)
        store, temp_dir, _ = fresh_store
        self._project(store)
        out = str(temp_dir / "t1.docx")
        services.auto_export_project("t1", out)
        from docx import Document
        texts = [p.text for p in Document(out).paragraphs]
        assert not any("интервью_ф13" in t for t in texts)
