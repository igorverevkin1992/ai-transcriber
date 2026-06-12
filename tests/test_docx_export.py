import re
import zipfile
from datetime import datetime, timezone

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Mm, Pt

from backend.docx_export import (
    _REFERENCE_PART_ORDER,
    generate_docx,
    is_legend_excluded_name,
)


class TestGenerateDocx:
    def test_header_is_download_name_with_docx_extension(self, tmp_path, sample_project):
        out = tmp_path / "out.docx"
        generate_docx(
            sample_project,
            final_map={"0": "Денис Майданов", "1": "Григорий Антипенко"},
            abbr_map={"0": "М", "1": "А"},
            output_path=str(out),
        )
        doc = Document(str(out))
        first_para = doc.paragraphs[0]
        assert first_para.text == "test_interview.docx"
        for run in first_para.runs:
            assert not run.bold
        assert first_para.alignment == WD_ALIGN_PARAGRAPH.JUSTIFY

    def test_legend_only_for_speakers_in_segments(self, tmp_path, sample_project):
        sample_project["result"]["speakers"]["99"] = {
            "duration_sec": 5.0, "suggested_name": "Призрак",
        }
        out = tmp_path / "out.docx"
        generate_docx(
            sample_project,
            final_map={"0": "Денис", "1": "Григорий", "99": "Призрак"},
            abbr_map={"0": "М", "1": "А", "99": "П"},
            output_path=str(out),
        )
        doc = Document(str(out))
        text = "\n".join(p.text for p in doc.paragraphs)
        assert "Призрак" not in text
        assert "Денис" in text
        assert "Григорий" in text

    def test_legend_ordered_by_first_appearance(self, tmp_path, sample_project):
        # Эталон ф5: ИНТЕРВЬЮЕР в легенде первый, хотя говорит меньше.
        # speakers_info отсортирован по длительности ("0" самый говорливый),
        # но первым в стенограмме появляется "1" — он и в легенде первый.
        sample_project["result"]["segments"] = [
            {"timecode": "11:04:15:00", "speaker": "1", "text": "Первый вопрос?"},
            {"timecode": "11:04:25:00", "speaker": "0", "text": "Длинный ответ."},
        ]
        out = tmp_path / "out.docx"
        generate_docx(
            sample_project,
            final_map={"0": "Майданов", "1": "Интервьюер"},
            abbr_map={"0": "М", "1": "И"},
            output_path=str(out),
        )
        doc = Document(str(out))
        texts = [p.text for p in doc.paragraphs]
        legend_lines = [t for t in texts if " – " in t and t.endswith(".")]
        assert legend_lines[0].startswith("Интервьюер")
        assert legend_lines[1].startswith("Майданов")

    def test_legend_exclude_hides_speaker_from_legend(self, tmp_path, sample_project):
        sample_project["result"]["speakers"]["2"] = {
            "duration_sec": 10.0, "suggested_name": "АЗК",
        }
        sample_project["result"]["segments"].append(
            {"timecode": "11:04:30:00", "speaker": "2", "text": "Ещё раз."}
        )
        out = tmp_path / "out.docx"
        generate_docx(
            sample_project,
            final_map={"0": "Денис", "1": "Григорий", "2": "АЗК"},
            abbr_map={"0": "М", "1": "А", "2": "АЗК"},
            output_path=str(out),
            legend_exclude={"2"},
        )
        doc = Document(str(out))
        legend_texts = [p.text for p in doc.paragraphs if "–" in p.text]
        assert not any("АЗК" in t for t in legend_texts)
        segment_texts = [p.text for p in doc.paragraphs if "11:04:30:00" in p.text]
        assert len(segment_texts) == 1
        assert "АЗК:" in segment_texts[0]

    def test_parenthetical_segment_no_speaker_prefix(self, tmp_path, sample_project):
        out = tmp_path / "out.docx"
        generate_docx(sample_project, {"0": "Д", "1": "Г"}, {"0": "М", "1": "А"}, str(out))
        doc = Document(str(out))
        tech_paras = [p for p in doc.paragraphs if "(Технические моменты.)" in p.text]
        assert len(tech_paras) == 1
        assert "М:" not in tech_paras[0].text
        assert tech_paras[0].text.startswith("11:04:15:00 (")

    def test_parenthetical_is_italic(self, tmp_path, sample_project):
        out = tmp_path / "out.docx"
        generate_docx(sample_project, {"0": "Д", "1": "Г"}, {"0": "М", "1": "А"}, str(out))
        doc = Document(str(out))
        found_italic = False
        for para in doc.paragraphs:
            for run in para.runs:
                if run.italic and "(" in run.text:
                    found_italic = True
        assert found_italic

    def test_no_page_numbers_in_footer(self, tmp_path, sample_project):
        out = tmp_path / "out.docx"
        generate_docx(sample_project, {"0": "Д", "1": "Г"}, {"0": "М", "1": "А"}, str(out))
        doc = Document(str(out))
        for section in doc.sections:
            footer_text = "".join(p.text for p in section.footer.paragraphs).strip()
            assert footer_text == ""

    def test_margins(self, tmp_path, sample_project):
        out = tmp_path / "out.docx"
        generate_docx(sample_project, {"0": "Д", "1": "Г"}, {"0": "М", "1": "А"}, str(out))
        doc = Document(str(out))
        section = doc.sections[0]
        assert abs(section.top_margin - Cm(2)) < 1000
        assert abs(section.left_margin - Cm(3)) < 1000
        assert abs(section.right_margin - Cm(1.5)) < 1000

    def test_header_footer_distance(self, tmp_path, sample_project):
        # Эталоны: header/footer distance 449580 EMU (~1.25 см).
        out = tmp_path / "out.docx"
        generate_docx(sample_project, {"0": "Д", "1": "Г"}, {"0": "М", "1": "А"}, str(out))
        doc = Document(str(out))
        section = doc.sections[0]
        assert section.header_distance == 449580
        assert section.footer_distance == 449580

    def test_trailing_empty_paragraph(self, tmp_path, sample_project):
        # Эталоны заканчиваются пустым абзацем после последней реплики.
        out = tmp_path / "out.docx"
        generate_docx(sample_project, {"0": "Д", "1": "Г"}, {"0": "М", "1": "А"}, str(out))
        doc = Document(str(out))
        assert not doc.paragraphs[-1].text.strip()

    def test_segment_format_inline_timecode(self, tmp_path, sample_project):
        out = tmp_path / "out.docx"
        generate_docx(sample_project, {"0": "Д", "1": "Г"}, {"0": "М", "1": "А"}, str(out))
        doc = Document(str(out))
        segment_paragraphs = [p for p in doc.paragraphs if "11:04:" in p.text]
        assert len(segment_paragraphs) == 3
        for para in segment_paragraphs:
            for run in para.runs:
                if run.text.startswith("11:04:"):
                    assert run.font.size == Pt(14) or run.font.size is None

    def test_returns_correct_download_name(self, tmp_path, sample_project):
        out = tmp_path / "out.docx"
        name = generate_docx(sample_project, {"0": "Д"}, {"0": "М"}, str(out))
        assert name == "test_interview.docx"

    def test_page_size_is_a4(self, tmp_path, sample_project):
        out = tmp_path / "out.docx"
        generate_docx(sample_project, {"0": "Д", "1": "Г"}, {"0": "М", "1": "А"}, str(out))
        doc = Document(str(out))
        section = doc.sections[0]
        assert abs(section.page_width - Mm(210)) < 1000
        assert abs(section.page_height - Mm(297)) < 1000

    def test_space_before_not_forced_to_zero(self, tmp_path, sample_project):
        out = tmp_path / "out.docx"
        generate_docx(sample_project, {"0": "Д", "1": "Г"}, {"0": "М", "1": "А"}, str(out))
        doc = Document(str(out))
        for para in doc.paragraphs:
            assert para.paragraph_format.space_before is None

    def test_normal_style_keeps_11pt_default(self, tmp_path, sample_project):
        # Human reference keeps the 11pt doc default (blank separator lines are
        # 11pt); 14pt is applied only per run.
        out = tmp_path / "out.docx"
        generate_docx(sample_project, {"0": "Д", "1": "Г"}, {"0": "М", "1": "А"}, str(out))
        doc = Document(str(out))
        assert doc.styles["Normal"].font.size is None

    def test_body_font_is_calibri(self, tmp_path, sample_project):
        # Reference body font is Calibri (theme minor font). The template's
        # Normal style carries no explicit rPr — the font resolves through the
        # theme, exactly as in the human reference.
        out = tmp_path / "out.docx"
        generate_docx(sample_project, {"0": "Д", "1": "Г"}, {"0": "М", "1": "А"}, str(out))
        with zipfile.ZipFile(str(out)) as z:
            theme = z.read("word/theme/theme1.xml").decode("utf-8")
        m = re.search(r'<a:minorFont>.*?<a:latin typeface="([^"]*)"', theme, re.S)
        assert m is not None
        assert m.group(1) == "Calibri"

    def test_parenthetical_trailing_period_is_italic(self, tmp_path, sample_project):
        # Reference italicizes the whole "(...)." including the trailing dot.
        sample_project["result"]["segments"] = [
            {"timecode": "11:04:15:00", "speaker": "0", "text": "(Технические моменты)."},
        ]
        out = tmp_path / "out.docx"
        generate_docx(sample_project, {"0": "Д"}, {"0": "М"}, str(out))
        doc = Document(str(out))
        para = next(p for p in doc.paragraphs if "Технические" in p.text)
        italic_run = next(r for r in para.runs if "Технические" in r.text)
        assert italic_run.italic
        assert italic_run.text.endswith(").")

    def test_nested_parenthetical_fully_italic(self, tmp_path, sample_project):
        sample_project["result"]["segments"] = [
            {"timecode": "11:04:15:00", "speaker": "0",
             "text": "(Технические моменты (перерыв).)"},
        ]
        out = tmp_path / "out.docx"
        generate_docx(sample_project, {"0": "Д"}, {"0": "М"}, str(out))
        doc = Document(str(out))
        paren_paras = [p for p in doc.paragraphs if "Технические моменты" in p.text]
        assert len(paren_paras) == 1
        paren_runs = [r for r in paren_paras[0].runs if "Технические" in r.text]
        assert len(paren_runs) == 1
        assert paren_runs[0].italic

    def test_unclear_standalone_has_speaker_prefix(self, tmp_path, sample_project):
        sample_project["result"]["segments"] = [
            {"timecode": "11:04:15:00", "speaker": "0", "text": "(неразборчиво)."},
        ]
        out = tmp_path / "out.docx"
        generate_docx(sample_project, {"0": "Д"}, {"0": "М"}, str(out))
        doc = Document(str(out))
        para = next(p for p in doc.paragraphs if "неразборчиво" in p.text)
        assert "М:" in para.text

    def test_unclear_not_italic(self, tmp_path, sample_project):
        sample_project["result"]["segments"] = [
            {"timecode": "11:04:15:00", "speaker": "0", "text": "Текст (неразборчиво) продолжение."},
        ]
        out = tmp_path / "out.docx"
        generate_docx(sample_project, {"0": "Д"}, {"0": "М"}, str(out))
        doc = Document(str(out))
        para = next(p for p in doc.paragraphs if "неразборчиво" in p.text)
        unclear_runs = [r for r in para.runs if "неразборчиво" in r.text]
        assert len(unclear_runs) == 1
        assert not unclear_runs[0].italic

    def test_segments_override_replaces_original(self, tmp_path, sample_project):
        edited = [
            {"timecode": "11:04:15:00", "speaker": "0", "text": "Исправленный текст."},
        ]
        out = tmp_path / "out.docx"
        generate_docx(
            sample_project,
            {"0": "Д", "1": "Г"}, {"0": "М", "1": "А"},
            str(out), segments_override=edited,
        )
        doc = Document(str(out))
        all_text = "\n".join(p.text for p in doc.paragraphs)
        assert "Исправленный текст." in all_text
        assert "Привет всем." not in all_text
        assert "Хорошо, давайте." not in all_text


class TestDocxStructure:
    def test_page_number_in_header(self, tmp_path, sample_project):
        out = tmp_path / "out.docx"
        generate_docx(sample_project, {"0": "Д", "1": "Г"}, {"0": "М", "1": "А"}, str(out))
        with zipfile.ZipFile(str(out)) as z:
            for name in z.namelist():
                if name.startswith("word/header"):
                    xml = z.read(name).decode("utf-8")
                    assert "PAGE" in xml
                    assert "MERGEFORMAT" in xml
                    break
            else:
                raise AssertionError("No header XML found")

    def test_page_number_header_right_aligned(self, tmp_path, sample_project):
        # The reference wraps the PAGE field in a w:sdt (page-number gallery),
        # so the field paragraph is not a direct child of w:hdr and is not
        # visible via section.header.paragraphs — inspect the header XML.
        out = tmp_path / "out.docx"
        generate_docx(sample_project, {"0": "Д", "1": "Г"}, {"0": "М", "1": "А"}, str(out))
        with zipfile.ZipFile(str(out)) as z:
            hdr = z.read("word/header1.xml").decode("utf-8")
        sdt = re.search(r"<w:sdt>.*?</w:sdt>", hdr, re.S)
        assert sdt is not None
        assert 'w:val="right"' in sdt.group(0)
        assert "PAGE" in sdt.group(0)

    def test_document_language_ru(self, tmp_path, sample_project):
        out = tmp_path / "out.docx"
        generate_docx(sample_project, {"0": "Д", "1": "Г"}, {"0": "М", "1": "А"}, str(out))
        with zipfile.ZipFile(str(out)) as z:
            styles_xml = z.read("word/styles.xml").decode("utf-8")
        assert 'w:val="ru-RU"' in styles_xml

    def test_blank_paragraph_mark_is_14pt(self, tmp_path, sample_project):
        # Reference paragraph marks (incl. blank separators) carry sz=28 so
        # blank lines render at 14pt, not the 11pt document default.
        out = tmp_path / "out.docx"
        generate_docx(sample_project, {"0": "Д", "1": "Г"}, {"0": "М", "1": "А"}, str(out))
        doc = Document(str(out))
        blank = next(p for p in doc.paragraphs if not p.text.strip())
        rpr = blank._p.get_or_add_pPr().find(qn("w:rPr"))
        assert rpr is not None
        sz = rpr.find(qn("w:sz"))
        assert sz is not None and sz.get(qn("w:val")) == "28"
        rfonts = rpr.find(qn("w:rFonts"))
        assert rfonts is not None and rfonts.get(qn("w:cstheme")) == "minorHAnsi"

    def test_runs_have_szcs(self, tmp_path, sample_project):
        out = tmp_path / "out.docx"
        generate_docx(sample_project, {"0": "Д", "1": "Г"}, {"0": "М", "1": "А"}, str(out))
        doc = Document(str(out))
        segment_para = next(p for p in doc.paragraphs if "11:04:" in p.text)
        run = segment_para.runs[0]
        szcs = run._element.find(qn("w:rPr")).find(qn("w:szCs"))
        assert szcs is not None
        assert szcs.get(qn("w:val")) == "28"

    def test_cols_space_708(self, tmp_path, sample_project):
        out = tmp_path / "out.docx"
        generate_docx(sample_project, {"0": "Д", "1": "Г"}, {"0": "М", "1": "А"}, str(out))
        doc = Document(str(out))
        cols = doc.sections[0]._sectPr.find(qn("w:cols"))
        assert cols.get(qn("w:space")) == "708"


    def test_header_has_trailing_empty_paragraph(self, tmp_path, sample_project):
        # Reference header = sdt-wrapped PAGE paragraph + one trailing empty
        # paragraph (2 w:p total in word/header1.xml).
        out = tmp_path / "out.docx"
        generate_docx(sample_project, {"0": "Д", "1": "Г"}, {"0": "М", "1": "А"}, str(out))
        with zipfile.ZipFile(str(out)) as z:
            hdr = z.read("word/header1.xml").decode("utf-8")
        assert len(re.findall(r"<w:p[ >]", hdr)) == 2

    def test_author_not_python_docx(self, tmp_path, sample_project):
        out = tmp_path / "out.docx"
        generate_docx(sample_project, {"0": "Д", "1": "Г"}, {"0": "М", "1": "А"}, str(out))
        doc = Document(str(out))
        assert doc.core_properties.author != "python-docx"

    def test_no_proof_state_in_settings(self, tmp_path, sample_project):
        out = tmp_path / "out.docx"
        generate_docx(sample_project, {"0": "Д", "1": "Г"}, {"0": "М", "1": "А"}, str(out))
        doc = Document(str(out))
        proof_state = doc.settings.element.find(qn("w:proofState"))
        assert proof_state is None

    def test_runs_have_cstheme(self, tmp_path, sample_project):
        out = tmp_path / "out.docx"
        generate_docx(sample_project, {"0": "Д", "1": "Г"}, {"0": "М", "1": "А"}, str(out))
        doc = Document(str(out))
        segment_para = next(p for p in doc.paragraphs if "11:04:" in p.text)
        run = segment_para.runs[0]
        rfonts = run._element.find(qn("w:rPr")).find(qn("w:rFonts"))
        assert rfonts is not None
        assert rfonts.get(qn("w:cstheme")) == "minorHAnsi"


class TestRevisionIds:
    """rsid/paraId-«шум», как у реально редактировавшихся в Word эталонов."""

    def test_every_body_paragraph_has_rsid_and_paraid(self, tmp_path, sample_project):
        out = tmp_path / "out.docx"
        generate_docx(sample_project, {"0": "Д", "1": "Г"}, {"0": "М", "1": "А"}, str(out))
        with zipfile.ZipFile(str(out)) as z:
            doc = z.read("word/document.xml").decode("utf-8")
        opens = re.findall(r"<w:p\b[^>]*>", doc)
        assert len(opens) > 0
        for tag in opens:
            assert "w:rsidR=" in tag
            assert "w:rsidRDefault=" in tag
            assert "w14:paraId=" in tag

    def test_settings_rsids_table_populated(self, tmp_path, sample_project):
        out = tmp_path / "out.docx"
        generate_docx(sample_project, {"0": "Д", "1": "Г"}, {"0": "М", "1": "А"}, str(out))
        with zipfile.ZipFile(str(out)) as z:
            settings = z.read("word/settings.xml").decode("utf-8")
        table = re.search(r"<w:rsids>.*?</w:rsids>", settings, re.S).group(0)
        assert len(re.findall(r"<w:rsid ", table)) >= 25

    def test_rsids_deterministic_for_same_filename(self, tmp_path, sample_project):
        a, b = tmp_path / "a.docx", tmp_path / "b.docx"
        generate_docx(sample_project, {"0": "Д", "1": "Г"}, {"0": "М", "1": "А"}, str(a))
        generate_docx(sample_project, {"0": "Д", "1": "Г"}, {"0": "М", "1": "А"}, str(b))
        with zipfile.ZipFile(str(a)) as z:
            da = z.read("word/document.xml")
        with zipfile.ZipFile(str(b)) as z:
            db = z.read("word/document.xml")
        assert da == db

    def test_paraids_high_bit_zero(self, tmp_path, sample_project):
        out = tmp_path / "out.docx"
        generate_docx(sample_project, {"0": "Д", "1": "Г"}, {"0": "М", "1": "А"}, str(out))
        with zipfile.ZipFile(str(out)) as z:
            doc = z.read("word/document.xml").decode("utf-8")
        for pid in re.findall(r'w14:paraId="([0-9A-F]{8})"', doc):
            assert pid[0] in "01234567"

    def test_runs_carry_rsids(self, tmp_path, sample_project):
        # Эталоны несут rsid на ~97% runs; раунд 11 ставил их только на абзацы.
        out = tmp_path / "out.docx"
        generate_docx(sample_project, {"0": "Д", "1": "Г"}, {"0": "М", "1": "А"}, str(out))
        with zipfile.ZipFile(str(out)) as z:
            doc = z.read("word/document.xml").decode("utf-8")
        runs = re.findall(r"<w:r(?: [^>]*)?>", doc)
        with_rsid = [r for r in runs if "w:rsid" in r]
        assert len(runs) > 0
        assert len(with_rsid) / len(runs) >= 0.9

    def test_paragraphs_carry_rsidrpr(self, tmp_path, sample_project):
        out = tmp_path / "out.docx"
        generate_docx(sample_project, {"0": "Д", "1": "Г"}, {"0": "М", "1": "А"}, str(out))
        with zipfile.ZipFile(str(out)) as z:
            doc = z.read("word/document.xml").decode("utf-8")
        for tag in re.findall(r"<w:p\b[^>]*>", doc):
            assert "w:rsidRPr=" in tag


class TestCapsFormatting:
    """Заголовочный блок эталонов (имя файла, разделитель, легенда) несёт w:caps."""

    def _generate(self, tmp_path, sample_project):
        out = tmp_path / "out.docx"
        generate_docx(sample_project, {"0": "Д", "1": "Г"}, {"0": "М", "1": "А"}, str(out))
        return Document(str(out))

    def _pmark_has_caps(self, para):
        rpr = para._p.get_or_add_pPr().find(qn("w:rPr"))
        return rpr is not None and rpr.find(qn("w:caps")) is not None

    def _run_has_caps(self, run):
        rpr = run._element.find(qn("w:rPr"))
        return rpr is not None and rpr.find(qn("w:caps")) is not None

    def test_header_paragraph_has_caps(self, tmp_path, sample_project):
        doc = self._generate(tmp_path, sample_project)
        header = doc.paragraphs[0]
        assert self._pmark_has_caps(header)
        assert self._run_has_caps(header.runs[0])

    def test_separator_after_header_has_caps_in_pmark(self, tmp_path, sample_project):
        doc = self._generate(tmp_path, sample_project)
        sep = doc.paragraphs[1]
        assert not sep.runs
        assert self._pmark_has_caps(sep)

    def test_legend_has_caps(self, tmp_path, sample_project):
        doc = self._generate(tmp_path, sample_project)
        legend = [p for p in doc.paragraphs if " – " in p.text]
        assert legend
        for para in legend:
            assert self._pmark_has_caps(para)
            assert all(self._run_has_caps(r) for r in para.runs)

    def test_content_and_second_separator_without_caps(self, tmp_path, sample_project):
        doc = self._generate(tmp_path, sample_project)
        # Разделитель после легенды (перед первой репликой) и реплики — без caps.
        for para in doc.paragraphs:
            if "11:04:" in para.text:
                assert not self._pmark_has_caps(para)
                assert not any(self._run_has_caps(r) for r in para.runs)
        # последний абзац — завершающий пустой, тоже без caps
        assert not self._pmark_has_caps(doc.paragraphs[-1])

    def test_caps_sits_between_rfonts_and_sz(self, tmp_path, sample_project):
        # Порядок элементов rPr как в эталонах: rFonts, caps, sz, szCs.
        out_path = tmp_path / "out.docx"
        generate_docx(sample_project, {"0": "Д", "1": "Г"}, {"0": "М", "1": "А"}, str(out_path))
        with zipfile.ZipFile(str(out_path)) as z:
            xml = z.read("word/document.xml").decode("utf-8")
        assert '<w:rFonts w:cstheme="minorHAnsi"/><w:caps/><w:sz w:val="28"/>' in xml


class TestRunSplitPattern:
    """Разбиение реплик на runs повторяет статистику эталонов (1/3/2-run)."""

    def _many_segments_project(self, n=120):
        segs = [
            {"timecode": f"12:{i // 60:02d}:{i % 60:02d}:00", "speaker": str(i % 2),
             "text": f"Реплика номер {i}, обычный текст без скобок."}
            for i in range(n)
        ]
        return {
            "original_filename": "Статистика_f1.mp4",
            "result": {
                "speakers": {
                    "0": {"duration_sec": 100.0, "suggested_name": "Тестов"},
                    "1": {"duration_sec": 50.0, "suggested_name": "Гость"},
                },
                "segments": segs,
            },
        }

    def test_run_count_distribution(self, tmp_path):
        out = tmp_path / "out.docx"
        generate_docx(self._many_segments_project(), {"0": "Т", "1": "Г"},
                      {"0": "Т", "1": "Г"}, str(out))
        doc = Document(str(out))
        turn_paras = [p for p in doc.paragraphs if p.text.startswith("12:")]
        assert len(turn_paras) == 120
        counts = {1: 0, 2: 0, 3: 0}
        for p in turn_paras:
            counts[len(p.runs)] = counts.get(len(p.runs), 0) + 1
        # Цензус f7/f8: ~35% 1-run, ~45% 3-run, ~20% 2-run (допуск ±15 п.п.)
        assert 0.20 <= counts[1] / 120 <= 0.50
        assert 0.30 <= counts[3] / 120 <= 0.60
        assert 0.05 <= counts[2] / 120 <= 0.35

    def test_three_run_split_at_frames_boundary(self, tmp_path):
        out = tmp_path / "out.docx"
        generate_docx(self._many_segments_project(), {"0": "Т", "1": "Г"},
                      {"0": "Т", "1": "Г"}, str(out))
        doc = Document(str(out))
        three_run = [p for p in doc.paragraphs
                     if p.text.startswith("12:") and len(p.runs) == 3]
        assert three_run
        for p in three_run:
            # run1 = «HH:MM:SS», run2 = «:FF » (с пробелом), run3 = «спикер: текст»
            assert re.fullmatch(r"\d{2}:\d{2}:\d{2}", p.runs[0].text)
            assert re.fullmatch(r":\d{2} ", p.runs[1].text)
            assert ": " in p.runs[2].text

    def test_frames_run_preserves_space(self, tmp_path):
        out = tmp_path / "out.docx"
        generate_docx(self._many_segments_project(), {"0": "Т", "1": "Г"},
                      {"0": "Т", "1": "Г"}, str(out))
        with zipfile.ZipFile(str(out)) as z:
            xml = z.read("word/document.xml").decode("utf-8")
        assert '<w:t xml:space="preserve">:00 </w:t>' in xml

    def test_one_run_contains_full_turn(self, tmp_path):
        out = tmp_path / "out.docx"
        generate_docx(self._many_segments_project(), {"0": "Т", "1": "Г"},
                      {"0": "Т", "1": "Г"}, str(out))
        doc = Document(str(out))
        one_run = [p for p in doc.paragraphs
                   if p.text.startswith("12:") and len(p.runs) == 1]
        assert one_run
        for p in one_run:
            assert re.match(r"\d{2}:\d{2}:\d{2}:\d{2} [ТГ]: ", p.runs[0].text)

    def test_parenthetical_segments_never_one_run(self, tmp_path, sample_project):
        # Курсивная скобочная ремарка не может жить в одном run-е с речью.
        sample_project["result"]["segments"] = [
            {"timecode": f"11:04:{i:02d}:00", "speaker": "0",
             "text": "Текст (смеётся) продолжение."}
            for i in range(30)
        ]
        out = tmp_path / "out.docx"
        generate_docx(sample_project, {"0": "Д"}, {"0": "М"}, str(out))
        doc = Document(str(out))
        for p in doc.paragraphs:
            if p.text.startswith("11:04:"):
                assert len(p.runs) >= 3
                italic_runs = [r for r in p.runs if r.italic]
                assert len(italic_runs) == 1
                assert italic_runs[0].text == "(смеётся)"


class TestGoBackBookmark:
    """Закладка _GoBack — артефакт курсора Word, есть во всех эталонах."""

    def test_single_goback_at_end_of_first_paragraph(self, tmp_path, sample_project):
        out = tmp_path / "out.docx"
        generate_docx(sample_project, {"0": "Д", "1": "Г"}, {"0": "М", "1": "А"}, str(out))
        with zipfile.ZipFile(str(out)) as z:
            doc = z.read("word/document.xml").decode("utf-8")
        assert doc.count('w:name="_GoBack"') == 1
        assert doc.count("<w:bookmarkEnd") == 1
        first_para = re.search(r"<w:p\b.*?</w:p>", doc, re.S).group(0)
        assert "_GoBack" in first_para
        assert first_para.rstrip().endswith('<w:bookmarkEnd w:id="0"/></w:p>')


class TestPageBreakMarkers:
    """lastRenderedPageBreak — Word помечает фактические разрывы страниц."""

    def _long_project(self):
        segs = [
            {"timecode": f"11:{i:02d}:00:00", "speaker": str(i % 2),
             "text": ("слово " * 40).strip() + "."}
            for i in range(20)
        ]
        return {
            "original_filename": "Длинный_f9.mp4",
            "result": {
                "speakers": {
                    "0": {"duration_sec": 100.0, "suggested_name": "Тестов"},
                    "1": {"duration_sec": 50.0, "suggested_name": "Гость"},
                },
                "segments": segs,
            },
        }

    def test_markers_equal_pages_minus_one(self, tmp_path):
        out = tmp_path / "out.docx"
        generate_docx(self._long_project(), {"0": "Т", "1": "Г"}, {"0": "Т", "1": "Г"}, str(out))
        with zipfile.ZipFile(str(out)) as z:
            doc = z.read("word/document.xml").decode("utf-8")
            app = z.read("docProps/app.xml").decode("utf-8")
        pages = int(re.search(r"<Pages>(\d+)</Pages>", app).group(1))
        assert pages > 1
        assert doc.count("lastRenderedPageBreak") == pages - 1

    def test_short_document_has_no_markers(self, tmp_path, sample_project):
        out = tmp_path / "out.docx"
        generate_docx(sample_project, {"0": "Д", "1": "Г"}, {"0": "М", "1": "А"}, str(out))
        with zipfile.ZipFile(str(out)) as z:
            doc = z.read("word/document.xml").decode("utf-8")
        assert "lastRenderedPageBreak" not in doc

    def test_marker_sits_in_run_after_rpr(self, tmp_path):
        out = tmp_path / "out.docx"
        generate_docx(self._long_project(), {"0": "Т", "1": "Г"}, {"0": "Т", "1": "Г"}, str(out))
        with zipfile.ZipFile(str(out)) as z:
            doc = z.read("word/document.xml").decode("utf-8")
        # маркер стоит сразу после </w:rPr> и перед <w:t> внутри run-а
        assert re.search(r"</w:rPr><w:lastRenderedPageBreak/><w:t", doc)


class TestDocumentMetadata:
    """app.xml/core.xml не выдают машинную генерацию (реальная статистика, даты)."""

    def test_app_word_count_matches_body_text(self, tmp_path, sample_project):
        out = tmp_path / "out.docx"
        generate_docx(sample_project, {"0": "Д", "1": "Г"}, {"0": "М", "1": "А"}, str(out))
        with zipfile.ZipFile(str(out)) as z:
            app = z.read("docProps/app.xml").decode("utf-8")
            doc = z.read("word/document.xml").decode("utf-8")
        body_text = "".join(re.findall(r"<w:t[^>]*>([^<]*)</w:t>", doc))
        words = int(re.search(r"<Words>(\d+)</Words>", app).group(1))
        assert words > 0
        assert words == len(body_text.split())

    def test_app_application_is_word(self, tmp_path, sample_project):
        out = tmp_path / "out.docx"
        generate_docx(sample_project, {"0": "Д", "1": "Г"}, {"0": "М", "1": "А"}, str(out))
        with zipfile.ZipFile(str(out)) as z:
            app = z.read("docProps/app.xml").decode("utf-8")
        assert "<Application>Microsoft Office Word</Application>" in app

    def test_core_dates_are_current(self, tmp_path, sample_project):
        out = tmp_path / "out.docx"
        generate_docx(sample_project, {"0": "Д", "1": "Г"}, {"0": "М", "1": "А"}, str(out))
        doc = Document(str(out))
        cp = doc.core_properties
        assert cp.created.year == datetime.now(timezone.utc).year
        assert cp.modified >= cp.created

    def test_core_revision_in_plausible_range(self, tmp_path, sample_project):
        out = tmp_path / "out.docx"
        generate_docx(sample_project, {"0": "Д", "1": "Г"}, {"0": "М", "1": "А"}, str(out))
        doc = Document(str(out))
        assert 15 <= doc.core_properties.revision <= 40


class TestTemplateSkeleton:
    """Пакет наследует скелет эталона, а не дефолтного шаблона python-docx."""

    def _generate(self, tmp_path, sample_project):
        out = tmp_path / "out.docx"
        generate_docx(sample_project, {"0": "Д", "1": "Г"}, {"0": "М", "1": "А"}, str(out))
        return out

    def test_no_python_docx_fingerprint_anywhere(self, tmp_path, sample_project):
        out = self._generate(tmp_path, sample_project)
        with zipfile.ZipFile(str(out)) as z:
            blob = b"".join(z.read(n) for n in z.namelist())
        assert b"python-docx" not in blob

    def test_part_list_matches_reference(self, tmp_path, sample_project):
        out = self._generate(tmp_path, sample_project)
        with zipfile.ZipFile(str(out)) as z:
            parts = set(z.namelist())
        # Word-родные части эталона присутствуют...
        assert "word/footnotes.xml" in parts
        assert "word/endnotes.xml" in parts
        # ...а отпечатки дефолтного шаблона python-docx отсутствуют.
        assert "docProps/thumbnail.jpeg" not in parts
        assert "word/numbering.xml" not in parts
        assert "word/stylesWithEffects.xml" not in parts
        assert not any(p.startswith("customXml/") for p in parts)

    def test_styles_are_reference_skeleton(self, tmp_path, sample_project):
        out = self._generate(tmp_path, sample_project)
        with zipfile.ZipFile(str(out)) as z:
            styles = z.read("word/styles.xml").decode("utf-8")
        # Эталон: компактный styles.xml с авто-ID a..a6 (русский Word);
        # дефолт python-docx — 349 КБ с сотней англоязычных стилей.
        assert len(styles) < 40000
        assert 'w:styleId="a3"' in styles

    def test_settings_are_reference_skeleton(self, tmp_path, sample_project):
        out = self._generate(tmp_path, sample_project)
        with zipfile.ZipFile(str(out)) as z:
            settings = z.read("word/settings.xml").decode("utf-8")
        assert '<w:zoom w:percent="100"/>' in settings
        assert '<w:decimalSymbol w:val=","/>' in settings
        assert '<w:listSeparator w:val=";"/>' in settings
        assert '<w:defaultTabStop w:val="708"/>' in settings

    def test_header_uses_sdt_with_noproof(self, tmp_path, sample_project):
        out = self._generate(tmp_path, sample_project)
        with zipfile.ZipFile(str(out)) as z:
            hdr = z.read("word/header1.xml").decode("utf-8")
        assert "<w:sdt>" in hdr
        assert "<w:noProof/>" in hdr


class TestPackageBytes:
    """zip-контейнер собран как Word, а не Python (порядок/таймстампы/декларации)."""

    def test_part_order_matches_reference(self, tmp_path, sample_project):
        out = tmp_path / "out.docx"
        generate_docx(sample_project, {"0": "Д", "1": "Г"}, {"0": "М", "1": "А"}, str(out))
        with zipfile.ZipFile(str(out)) as z:
            order = [i.filename for i in z.infolist()]
        assert order == list(_REFERENCE_PART_ORDER)

    def test_all_timestamps_dos_epoch(self, tmp_path, sample_project):
        out = tmp_path / "out.docx"
        generate_docx(sample_project, {"0": "Д", "1": "Г"}, {"0": "М", "1": "А"}, str(out))
        with zipfile.ZipFile(str(out)) as z:
            assert all(i.date_time == (1980, 1, 1, 0, 0, 0) for i in z.infolist())

    def test_entries_are_windows_made(self, tmp_path, sample_project):
        out = tmp_path / "out.docx"
        generate_docx(sample_project, {"0": "Д", "1": "Г"}, {"0": "М", "1": "А"}, str(out))
        with zipfile.ZipFile(str(out)) as z:
            for i in z.infolist():
                assert i.create_system == 0
                assert i.external_attr == 0

    def test_xml_declaration_word_style(self, tmp_path, sample_project):
        out = tmp_path / "out.docx"
        generate_docx(sample_project, {"0": "Д", "1": "Г"}, {"0": "М", "1": "А"}, str(out))
        with zipfile.ZipFile(str(out)) as z:
            raw = z.read("word/document.xml")
        assert raw.startswith(b'<?xml version="1.0" encoding="UTF-8" standalone="yes"?>\r\n')

    def test_package_integrity_preserved(self, tmp_path, sample_project):
        out = tmp_path / "out.docx"
        generate_docx(sample_project, {"0": "Д", "1": "Г"}, {"0": "М", "1": "А"}, str(out))
        with zipfile.ZipFile(str(out)) as z:
            assert z.testzip() is None
        # python-docx по-прежнему открывает файл
        assert len(Document(str(out)).paragraphs) > 0


class TestLegendExcludedNames:
    def test_azk_gzk_excluded(self):
        assert is_legend_excluded_name("АЗК")
        assert is_legend_excluded_name("ГЗК")

    def test_m_zh_labels_excluded(self):
        assert is_legend_excluded_name("М")
        assert is_legend_excluded_name("Ж")
        assert is_legend_excluded_name("М1")
        assert is_legend_excluded_name("М2")
        assert is_legend_excluded_name("ММ")

    def test_zhenschina_muzhchina_excluded(self):
        assert is_legend_excluded_name("ЖЕНЩИНА")
        assert is_legend_excluded_name("МУЖЧИНА")
        assert is_legend_excluded_name("Женщина")
        assert is_legend_excluded_name("Мужчина")

    def test_named_speakers_not_excluded(self):
        assert not is_legend_excluded_name("Антипенко")
        assert not is_legend_excluded_name("Интервьюер")
        assert not is_legend_excluded_name("Денис")
