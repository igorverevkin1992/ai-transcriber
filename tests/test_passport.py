"""Тесты парсинга «паспорта съёмки» (.docx)."""

from docx import Document

from backend.passport import _split_names, parse_passport, read_passport_text


def _make_table_passport(path, rows):
    doc = Document()
    table = doc.add_table(rows=0, cols=2)
    for label, value in rows:
        cells = table.add_row().cells
        cells[0].text = label
        cells[1].text = value
    doc.save(str(path))


def _make_para_passport(path, lines):
    doc = Document()
    for line in lines:
        doc.add_paragraph(line)
    doc.save(str(path))


class TestSplitNames:
    def test_comma_separated(self):
        assert _split_names("Олег Александрович, Галина Васильевна") == [
            "Олег Александрович", "Галина Васильевна"]

    def test_semicolon_and_newline(self):
        assert _split_names("Иванов; Петров\nСидоров") == ["Иванов", "Петров", "Сидоров"]

    def test_numbered_list(self):
        assert _split_names("1) Иванов\n2) Петров") == ["Иванов", "Петров"]

    def test_empty(self):
        assert _split_names("") == []


class TestParsePassportTable:
    def test_full_form(self, tmp_path):
        p = tmp_path / "passport.docx"
        _make_table_passport(p, [
            ("Описание съёмки", "Тренировка по бадминтону в вузе"),
            ("Герои", "Олег Александрович, Галина Васильевна"),
            ("Количество героев", "2"),
        ])
        data = parse_passport(p)
        assert data["speakers"] == ["Олег Александрович", "Галина Васильевна"]
        assert data["num_heroes"] == 2
        assert "бадминтон" in data["description"].lower()

    def test_count_from_names_when_absent(self, tmp_path):
        p = tmp_path / "p.docx"
        _make_table_passport(p, [("Гости", "Иванов; Петров; Сидоров")])
        data = parse_passport(p)
        assert data["num_heroes"] == 3
        assert len(data["speakers"]) == 3


class TestParsePassportParagraphs:
    def test_label_value_lines(self, tmp_path):
        p = tmp_path / "p.docx"
        _make_para_passport(p, [
            "Описание: интервью о кино",
            "Герои: Майданов Денис",
            "Количество героев: 1",
        ])
        data = parse_passport(p)
        assert data["speakers"] == ["Майданов Денис"]
        assert data["num_heroes"] == 1
        assert data["description"] == "интервью о кино"


class TestRealTemplate:
    def test_value_under_label_and_crew(self, tmp_path):
        # Шаблон пользователя: подпись на строке, значение — ниже; группа отдельными
        # полями; «Герой» в ед. числе; «Что снято» как описание.
        p = tmp_path / "p.docx"
        _make_para_passport(p, [
            "Дата съемки: 17.06.2026",
            "Локация: вуз",
            "Герой:",
            "Олег Александрович",
            "Галина Васильевна",
            "",
            "Что снято:",
            "Тренировка по бадминтону, интервью с тренером.",
            "",
            "Автор: Михаил Ломов",
            "Оператор: Иван Петров",
            "Инженер: Семён Сидоров",
        ])
        data = parse_passport(p)
        assert data["speakers"] == ["Олег Александрович", "Галина Васильевна"]
        assert data["num_heroes"] == 2
        assert data["has_host"] is True
        assert data["crew"] == ["Иван Петров", "Семён Сидоров"]
        assert "бадминтон" in data["description"].lower()
        # Дата/Локация игнорируются, в группу/героев не попадают.
        assert "17.06.2026" not in data["description"]

    def test_inline_hero_singular(self, tmp_path):
        p = tmp_path / "p.docx"
        _make_para_passport(p, ["Герой: Майданов Денис", "Что снято: интервью о кино"])
        data = parse_passport(p)
        assert data["speakers"] == ["Майданов Денис"]
        assert data["description"] == "интервью о кино"

    def test_host_negative_disables(self, tmp_path):
        p = tmp_path / "p.docx"
        _make_para_passport(p, ["Герой: Иванов", "Ведущий: нет"])
        data = parse_passport(p)
        assert data["has_host"] is False

    def test_explicit_host_no_beats_author(self, tmp_path):
        # «Ведущий: нет» решает, даже если «Автор» заполнен именем.
        p = tmp_path / "p.docx"
        _make_para_passport(p, [
            "Герой: Иванов, Петров",
            "Ведущий: нет",
            "Автор: Михаил Ломов",
        ])
        data = parse_passport(p)
        assert data["has_host"] is False

    def test_host_negative_with_trailing_dot(self, tmp_path):
        p = tmp_path / "p.docx"
        _make_para_passport(p, ["Герой: Иванов", "Ведущий за кадром: Нет."])
        assert parse_passport(p)["has_host"] is False

    def test_author_alone_implies_host(self, tmp_path):
        # Поля «Ведущий» нет — заполненный «Автор» означает закадрового ведущего.
        p = tmp_path / "p.docx"
        _make_para_passport(p, ["Герой: Иванов", "Автор: Михаил Ломов"])
        assert parse_passport(p)["has_host"] is True

    def test_host_default_true_when_absent(self, tmp_path):
        p = tmp_path / "p.docx"
        _make_para_passport(p, ["Герой: Иванов"])
        data = parse_passport(p)
        assert data["has_host"] is True

    def test_crew_comma_separated(self, tmp_path):
        p = tmp_path / "p.docx"
        _make_para_passport(p, ["Герой: Иванов", "Съёмочная группа: Петров, Сидоров"])
        data = parse_passport(p)
        assert data["crew"] == ["Петров", "Сидоров"]


class TestParsePassportEdgeCases:
    def test_missing_file(self, tmp_path):
        assert parse_passport(tmp_path / "nope.docx") is None

    def test_empty_doc_returns_none(self, tmp_path):
        p = tmp_path / "empty.docx"
        Document().save(str(p))
        assert parse_passport(p) is None

    def test_read_passport_text_collects_all(self, tmp_path):
        p = tmp_path / "p.docx"
        _make_table_passport(p, [("Герои", "Иванов")])
        text = read_passport_text(p)
        assert "Герои" in text and "Иванов" in text


class TestParticipantsField:
    def _docx(self, tmp_path, lines):
        from docx import Document as Doc
        d = Doc()
        for line in lines:
            d.add_paragraph(line)
        p = tmp_path / "p.docx"
        d.save(str(p))
        return p

    def test_participants_with_roles(self, tmp_path):
        from backend.passport import parse_passport
        data = parse_passport(self._docx(tmp_path, [
            "Герой: Канаева Евгения",
            "Снимаются: Батыршина Яна (ведущая), Канаева Светлана",
            "Что снято: интервью с мамой героини",
        ]))
        assert data["speakers"] == ["Канаева Евгения"]
        assert data["participants"] == ["Батыршина Яна", "Канаева Светлана"]
        assert data["host_name"] == "Батыршина Яна"

    def test_v_kadre_alias_and_multiline(self, tmp_path):
        from backend.passport import parse_passport
        data = parse_passport(self._docx(tmp_path, [
            "В кадре:",
            "Батыршина Яна (ведущая)",
            "Канаева Светлана",
        ]))
        assert data["participants"] == ["Батыршина Яна", "Канаева Светлана"]
        assert data["host_name"] == "Батыршина Яна"

    def test_uchastniki_now_participants(self, tmp_path):
        from backend.passport import parse_passport
        data = parse_passport(self._docx(tmp_path, ["Участники: Иванов Иван"]))
        assert data["participants"] == ["Иванов Иван"]
        assert data["speakers"] == []

    def test_no_participants_heroes_still_speakers(self, tmp_path):
        from backend.passport import parse_passport
        data = parse_passport(self._docx(tmp_path, [
            "Герои: Ковальская Марина", "Количество героев: 1",
        ]))
        assert data["speakers"] == ["Ковальская Марина"]
        assert data["participants"] == []
        assert data["host_name"] is None
