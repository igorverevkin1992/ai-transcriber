"""Парсинг «паспорта съёмки» (.docx).

После каждой съёмки ассистент заполняет форму в Word. Достоверные поля подаются в
пайплайн как приоритетный источник `meta`:
- **Герои** (гости в кадре) → имена гостей в легенде, число → хинт диаризации;
- **Ведущий за кадром** → голосов = герои + 1 (ведущий помечается `АЗК`);
- **Съёмочная группа** (оператор/инженер/продюсер/ассистенты) → их закадровая речь
  помечается `(Технические моменты)`, а не приписывается гостю;
- **Что снято** (тема) → контекст для Gemini-правки границ.

Термины/глоссарий и стартовый таймкод в паспорт НЕ входят (глоссарий — отдельный
механизм; таймкод читается из контейнера `.wmv`).

Разбор: подписи полей (с двоеточием) + значение в той же строке ЛИБО на следующих
строках под подписью (до следующей подписи). Поддержаны абзацы и 2-кол. таблицы.
Если поля не распознаны — вызывающий код прибегает к Gemini-фолбэку
(`gemini_extract_passport`), передав сырой текст (`read_passport_text`).
"""

from pathlib import Path

from backend.config import logger

# Группы подписей полей (нормализованные к нижнему регистру). Порядок важен:
# «количество героев» проверяется раньше «герои». Матч по равенству/startswith.
_FIELD_LABELS = [
    ("count", ("количество героев", "число героев", "кол-во героев", "количество", "кол-во")),
    # Явное поле «Ведущий» — решающее. «Автор/корреспондент/интервьюер» — слабый
    # сигнал наличия ведущего: учитывается ТОЛЬКО если поля «Ведущий» нет
    # (иначе «Ведущий: нет» + «Автор: Имя» ошибочно давали has_host=True).
    ("host", ("ведущий",)),
    ("host_weak", ("автор", "корреспондент", "интервьюер")),
    ("crew", ("съёмочная группа", "съемочная группа", "группа", "оператор", "инженер",
              "продюсер", "ассистент", "режиссёр", "режиссер", "звукорежиссёр", "звукорежиссер")),
    ("heroes", ("имена героев", "герои", "герой", "гость", "гости", "участники", "спикеры")),
    ("description", ("что снято", "описание съёмки", "описание съемки", "описание",
                     "синопсис", "тема", "о чём", "о чем")),
    ("ignore", ("дата съёмки", "дата съемки", "дата", "локация", "место съёмки", "место")),
]

# Значение поля «Ведущий», означающее отсутствие закадрового ведущего.
_HOST_NEGATIVE = {"нет", "no", "-", "—", "–", "не было", "без ведущего", "отсутствует"}


def _is_host_negative(value: str) -> bool:
    return value.strip(" \t.!—–-").lower() in _HOST_NEGATIVE

# Разделители списка имён внутри одного значения.
_NAME_SPLIT = (";", ",", "\n", "•", " - ")


def _norm(s: str) -> str:
    return (s or "").strip().lower().rstrip(":").strip()


def _match_field(label_text: str) -> str | None:
    """Сопоставить подпись поля с ключом группы (или None)."""
    norm = _norm(label_text)
    if not norm:
        return None
    for key, labels in _FIELD_LABELS:
        for lbl in labels:
            if norm == lbl or norm.startswith(lbl):
                return key
    return None


def _split_names(value: str) -> list[str]:
    """Разбить значение поля на отдельные имена (по строкам/запятым/точкам с зпт)."""
    if not value:
        return []
    parts = [value]
    for sep in _NAME_SPLIT:
        parts = [p for chunk in parts for p in chunk.split(sep)]
    out = []
    for p in parts:
        name = p.strip(" \t.-–—•").strip()
        name = name.lstrip("0123456789).").strip()  # «1) Иванов», «2. Петров»
        if name and name not in out:
            out.append(name)
    return out


class _Fields:
    """Аккумулятор значений полей паспорта (поля группы/героев накапливаются)."""

    def __init__(self):
        self.heroes: list[str] = []
        self.crew: list[str] = []
        self.host_values: list[str] = []
        self.host_seen = False
        self.host_weak_values: list[str] = []
        self.count_text: str | None = None
        self.desc_lines: list[str] = []

    def add(self, field: str | None, value: str):
        value = (value or "").strip()
        if field == "heroes":
            self.heroes.extend(n for n in _split_names(value) if n not in self.heroes)
        elif field == "crew":
            self.crew.extend(n for n in _split_names(value) if n not in self.crew)
        elif field == "host":
            self.host_seen = True
            if value:
                self.host_values.append(value)
        elif field == "host_weak":
            if value:
                self.host_weak_values.append(value)
        elif field == "count":
            if self.count_text is None and value:
                self.count_text = value
        elif field == "description":
            if value:
                self.desc_lines.append(value)
        # field == "ignore" / None → пропускаем

    def to_dict(self) -> dict | None:
        description = " ".join(self.desc_lines).strip()
        num = 0
        if self.count_text:
            digits = "".join(ch for ch in self.count_text if ch.isdigit())
            if digits:
                num = int(digits)
        if num <= 0:
            num = len(self.heroes)

        if self.host_seen and self.host_values:
            # Явное поле «Ведущий» решает («нет» → False), «Автор» его не перебивает.
            has_host = not all(_is_host_negative(v) for v in self.host_values)
        elif self.host_weak_values:
            # Поля «Ведущий» нет — наличие заполненного «Автор/корреспондент»
            # трактуем как наличие закадрового ведущего.
            has_host = not all(_is_host_negative(v) for v in self.host_weak_values)
        else:
            has_host = True  # дефолт: интервью обычно с закадровым ведущим

        if not self.heroes and num <= 0 and not description and not self.crew:
            return None
        return {
            "speakers": self.heroes,
            "num_heroes": num,
            "has_host": has_host,
            "crew": self.crew,
            "description": description,
        }


def _parse_paragraphs(doc, fields: _Fields):
    """Разбор абзацев: подпись с двоеточием открывает поле, строки без двоеточия —
    продолжение значения под подписью."""
    current: str | None = None
    for para in doc.paragraphs:
        line = para.text.strip()
        if not line:
            continue
        if ":" in line:
            label, _, value = line.partition(":")
            field = _match_field(label)
            if field is not None:
                current = field
                fields.add(field, value)
                continue
            # Двоеточие, но не известная подпись → продолжение текущего поля.
            if current is not None:
                fields.add(current, line)
        elif current is not None:
            fields.add(current, line)  # значение под подписью


def _parse_tables(doc, fields: _Fields):
    """Разбор 2-кол. таблиц: ячейка-подпись | значение."""
    for table in getattr(doc, "tables", []):
        for row in table.rows:
            cells = [c.text for c in row.cells]
            if len(cells) >= 2 and cells[0].strip():
                fields.add(_match_field(cells[0]), " ".join(c for c in cells[1:] if c.strip()))


def read_passport_text(path) -> str:
    """Сырой текст паспорта (абзацы + ячейки таблиц) — для Gemini-фолбэка."""
    path = Path(path)
    if not path.exists():
        return ""
    try:
        from docx import Document
    except ImportError:
        return ""
    try:
        doc = Document(str(path))
    except Exception:
        return ""
    lines = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in getattr(doc, "tables", []):
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                lines.append(" | ".join(cells))
    return "\n".join(lines)


def parse_passport(path) -> dict | None:
    """Распарсить .docx-паспорт детерминированно по подписям формы.

    Возвращает `{"speakers": [имена героев], "num_heroes": int, "has_host": bool,
    "crew": [имена группы], "description": str}` либо `None`, если файл не открылся /
    поля не распознаны (тогда вызывающий код прибегает к Gemini-фолбэку).
    """
    path = Path(path)
    if not path.exists():
        return None
    try:
        from docx import Document
    except ImportError:
        logger.warning("python-docx недоступен — паспорт не распарсен")
        return None
    try:
        doc = Document(str(path))
    except Exception as e:
        logger.warning("Не удалось открыть паспорт %s: %s", path.name, e)
        return None

    fields = _Fields()
    _parse_tables(doc, fields)
    _parse_paragraphs(doc, fields)
    return fields.to_dict()
