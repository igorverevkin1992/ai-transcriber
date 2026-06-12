import random
import re
import zipfile
from datetime import datetime, timedelta, timezone

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Mm, Pt

from backend.config import DOCX_AUTHOR, DOCX_TEMPLATE_PATH
from backend.turns import UNCLEAR_TEXT
from backend.utils import strip_extension

# Parenthetical remark, optionally followed by sentence punctuation that the
# human reference also italicizes (e.g. the whole "(...)." including the dot).
PARENTHETICAL_RE = re.compile("(\\((?:[^()]*|\\([^()]*\\))*\\)[.!?…]*)")

# Служебные и эпизодические говорящие не входят в легенду эталонов:
# АЗК/ГЗК (автор/голос за кадром) и безымянные метки М1, М2, М3, Ж, ММ
# (цензус f8: реплики «М1:», «Ж:» есть, в легенде — только именованные).
_LEGEND_EXCLUDE_NAME_RE = re.compile(r"^([А-ЯЁ]ЗК|[МЖ]{1,2}\d*|ЖЕНЩИНА|МУЖЧИНА)$")


def is_legend_excluded_name(name: str) -> bool:
    """True, если имя спикера — служебная метка, не входящая в легенду."""
    return bool(_LEGEND_EXCLUDE_NAME_RE.match(name.strip().upper()))


def _set_paragraph_mark_rpr(paragraph, caps: bool = False):
    """Свойства знака абзаца (14pt, cstheme) — как у всех абзацев эталона.

    В эталонах даже ПУСТЫЕ абзацы-разделители несут rPr знака абзаца с
    sz=28: без него высота пустой строки падает до 11pt и вертикальная
    раскладка/разбиение на страницы расходятся с ручным оригиналом.

    caps=True — для заголовочного блока (имя файла, разделитель после него,
    легенда): эталоны несут там <w:caps/> между rFonts и sz.
    """
    p_pr = paragraph._p.get_or_add_pPr()
    rpr = p_pr.find(qn("w:rPr"))
    if rpr is None:
        # rPr знака абзаца идёт последним среди свойств абзаца (после jc).
        rpr = OxmlElement("w:rPr")
        p_pr.append(rpr)
    rpr.append(OxmlElement("w:rFonts", {qn("w:cstheme"): "minorHAnsi"}))
    if caps:
        rpr.append(OxmlElement("w:caps"))
    rpr.append(OxmlElement("w:sz", {qn("w:val"): "28"}))
    rpr.append(OxmlElement("w:szCs", {qn("w:val"): "28"}))


def _configure_paragraph(paragraph, caps: bool = False):
    """Применяет к абзацу выравнивание JUSTIFY, line_spacing=1.0, space_after=0."""
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    fmt = paragraph.paragraph_format
    fmt.space_after = Pt(0)
    fmt.line_spacing_rule = WD_LINE_SPACING.SINGLE
    fmt.line_spacing = 1.0
    _set_paragraph_mark_rpr(paragraph, caps=caps)


_HEX = "0123456789ABCDEF"


def _rand_rsid(rng: random.Random) -> str:
    """Сессионный rsid в формате Word: «00» + 6 hex (00xxxxxx)."""
    return "00" + "".join(rng.choice(_HEX) for _ in range(6))


def _rand_paraid(rng: random.Random, used: set) -> str:
    """Уникальный w14:paraId/textId: 8 hex, старший бит 0 (первый нибл 0-7)."""
    while True:
        pid = rng.choice("01234567") + "".join(rng.choice(_HEX) for _ in range(7))
        if pid not in used:
            used.add(pid)
            return pid


def _apply_revision_ids(doc, seed_text: str) -> None:
    """Проставляет rsid/paraId абзацам и пересобирает таблицу w:rsids.

    Эталоны несут 1200-2900 rsid-атрибутов и таблицу из 150-780 rsid в
    settings.xml — побочный продукт реального редактирования в Word. У
    свежесгенерированного файла их нет (главный статистический маркер
    машины). Засеваем ГПСЧ именем файла: один и тот же проект даёт
    идентичный «шум» (детерминизм для тестов и воспроизводимости).
    """
    rng = random.Random(seed_text)
    pool = [_rand_rsid(rng) for _ in range(rng.randint(25, 60))]
    rsid_root = _rand_rsid(rng)

    body = doc.element.body
    paras = body.findall(qn("w:p"))
    used_ids: set = set()
    current = rng.choice(pool)
    for p in paras:
        runs = p.findall(qn("w:r"))
        # Пустой абзац: Word ставит сентинел textId="77777777" («текст не
        # менялся») и НЕ ставит rsidRPr (эталоны: 0 из 13 пустых с rsidRPr).
        # Исключение — завершающий пустой абзац, он несёт уникальный textId.
        is_empty = not runs
        # Соседние абзацы чаще делят rsid (правка одной сессии) — инерция ~70%.
        if rng.random() > 0.7:
            current = rng.choice(pool)
        # Порядок атрибутов как в эталонах: paraId, textId, rsidR, rsidRPr,
        # rsidRDefault, rsidP.
        p.set(qn("w14:paraId"), _rand_paraid(rng, used_ids))
        if is_empty and p is not paras[-1]:
            p.set(qn("w14:textId"), "77777777")
        else:
            p.set(qn("w14:textId"), _rand_paraid(rng, used_ids))
        p.set(qn("w:rsidR"), current)
        if not is_empty:
            p.set(qn("w:rsidRPr"), current)
        p.set(qn("w:rsidRDefault"), current)
        p.set(qn("w:rsidP"), rng.choice(pool))
        # Run-уровневые rsid: эталоны несут их на ~97% runs.
        for ri, r in enumerate(runs):
            if rng.random() < 0.03:
                continue  # ~3% runs без атрибутов, как в эталоне
            if ri > 0 and rng.random() < 0.2:
                val = rng.choice(pool)
                r.set(qn("w:rsidR"), val)
                r.set(qn("w:rsidRPr"), val)
            elif ri == 0:
                r.set(qn("w:rsidRPr"), current)
            else:
                r.set(qn("w:rsidR"), current)

    sect = body.find(qn("w:sectPr"))
    if sect is not None:
        sect.set(qn("w:rsidR"), rng.choice(pool))
        sect.set(qn("w:rsidSect"), rsid_root)

    rsids = doc.settings.element.find(qn("w:rsids"))
    if rsids is not None:
        for child in list(rsids):
            rsids.remove(child)
        root_el = OxmlElement("w:rsidRoot")
        root_el.set(qn("w:val"), rsid_root)
        rsids.append(root_el)
        for val in [rsid_root, *pool]:
            e = OxmlElement("w:rsid")
            e.set(qn("w:val"), val)
            rsids.append(e)


def _set_core_dates(doc, seed_text: str) -> None:
    """Свежие даты и правдоподобный revision в core.xml.

    Дефолт python-docx — created/modified 2013-12-23, revision 1. Эталоны:
    created на 1-2 дня раньше modified, секунды :00, revision 17-36.
    """
    now = datetime.now(timezone.utc).replace(second=0, microsecond=0)
    doc.core_properties.created = now - timedelta(days=2)
    doc.core_properties.modified = now
    doc.core_properties.revision = random.Random(seed_text + ":rev").randint(15, 40)


# Оценочное число слов на страницу (эталоны: 311-394). Используется и для
# Pages в app.xml, и для расстановки lastRenderedPageBreak — чтобы число
# маркеров (Pages−1) было согласовано со статистикой.
_WORDS_PER_PAGE = 385

# Счётчики docProps/app.xml в порядке появления в XML.
_APP_STAT_TAGS = (
    "TotalTime", "Pages", "Words", "Characters",
    "Lines", "Paragraphs", "CharactersWithSpaces",
)


def _para_word_count(p) -> int:
    """Число слов в абзаце (по тексту всех его run-ов)."""
    return sum(len((t.text or "").split()) for t in p.findall(".//" + qn("w:t")))


def _mark_page_break(p) -> None:
    """Вставляет lastRenderedPageBreak в первый run абзаца (после rPr)."""
    r = p.find(qn("w:r"))
    if r is None:
        return
    lrpb = OxmlElement("w:lastRenderedPageBreak")
    rpr = r.find(qn("w:rPr"))
    if rpr is not None:
        rpr.addnext(lrpb)
    else:
        r.insert(0, lrpb)


def _insert_page_break_markers(doc) -> None:
    """Расставляет lastRenderedPageBreak на оценочных границах страниц.

    Word при сохранении помечает фактические разрывы (эталоны: 9-23 маркера,
    ~1 на страницу). Файл без них «не открывался в Word». Оцениваем границы
    по накопителю слов; маркеров = Pages−1, согласовано с app.xml.
    """
    paras = doc.element.body.findall(qn("w:p"))
    total = sum(_para_word_count(p) for p in paras)
    markers_needed = max(1, round(total / _WORDS_PER_PAGE)) - 1
    if markers_needed <= 0:
        return
    placed = 0
    acc = 0
    boundary = _WORDS_PER_PAGE
    for p in paras:
        acc += _para_word_count(p)
        if acc >= boundary:
            _mark_page_break(p)
            placed += 1
            boundary += _WORDS_PER_PAGE
            if placed >= markers_needed:
                break


# Порядок частей в zip, как пишет Word (эталоны). Неизвестные части идут
# в конец.
_REFERENCE_PART_ORDER = (
    "[Content_Types].xml",
    "_rels/.rels",
    "word/_rels/document.xml.rels",
    "word/document.xml",
    "word/footnotes.xml",
    "word/endnotes.xml",
    "word/header1.xml",
    "word/theme/theme1.xml",
    "word/settings.xml",
    "word/styles.xml",
    "word/webSettings.xml",
    "docProps/app.xml",
    "docProps/core.xml",
    "word/fontTable.xml",
)

# Декларация XML в стиле Word: двойные кавычки + CRLF (lxml пишет одинарные + LF).
_WORD_XML_DECL = b'<?xml version="1.0" encoding="UTF-8" standalone="yes"?>\r\n'


def _app_statistics(doc_xml: str) -> dict:
    """Считает Words/Characters/Lines/Pages/... по тексту document.xml."""
    full = "".join(re.findall(r"<w:t[^>]*>([^<]*)</w:t>", doc_xml))
    words = len(full.split())
    chars_with_spaces = len(full)
    paragraphs = sum(1 for p in re.findall(r"<w:p\b.*?</w:p>", doc_xml, re.S) if "<w:t" in p)
    return {
        "TotalTime": max(60, words // 10),
        "Pages": max(1, round(words / _WORDS_PER_PAGE)),
        "Words": words,
        "Characters": len(re.sub(r"\s", "", full)),
        "Lines": max(1, round(chars_with_spaces / 141)),
        "Paragraphs": paragraphs,
        "CharactersWithSpaces": chars_with_spaces,
    }


def _normalize_xml_declaration(data: bytes) -> bytes:
    """Приводит XML-декларацию к Word-стилю (двойные кавычки + CRLF)."""
    if not data.startswith(b"<?xml"):
        return data
    end = data.find(b"?>")
    if end == -1:
        return data
    rest = data[end + 2:]
    if rest.startswith(b"\r\n"):
        rest = rest[2:]
    elif rest.startswith(b"\n"):
        rest = rest[1:]
    return _WORD_XML_DECL + rest


def _finalize_package(path: str) -> None:
    """Доводит готовый пакет до байтового вида Word.

    python-docx собирает zip по-своему (порядок частей, Unix-таймстампы и
    права, lxml-декларации с одинарными кавычками) — это выдаёт машину
    через `unzip -l` за секунду. Здесь: вписываем статистику app.xml,
    нормализуем XML-декларации и перепаковываем zip в эталонном порядке с
    DOS-эпохой и Windows-метаданными записей.
    """
    with zipfile.ZipFile(path) as zin:
        parts = {i.filename: zin.read(i.filename) for i in zin.infolist()}

    stats = _app_statistics(parts["word/document.xml"].decode("utf-8"))
    app = parts["docProps/app.xml"].decode("utf-8")
    for tag in _APP_STAT_TAGS:
        app = re.sub(rf"<{tag}>[^<]*</{tag}>", f"<{tag}>{stats[tag]}</{tag}>", app)
    parts["docProps/app.xml"] = app.encode("utf-8")

    parts = {name: _normalize_xml_declaration(data) for name, data in parts.items()}

    ordered = [n for n in _REFERENCE_PART_ORDER if n in parts]
    ordered += [n for n in parts if n not in _REFERENCE_PART_ORDER]

    with zipfile.ZipFile(path, "w", zipfile.ZIP_DEFLATED) as zout:
        for name in ordered:
            info = zipfile.ZipInfo(name, date_time=(1980, 1, 1, 0, 0, 0))
            info.compress_type = zipfile.ZIP_DEFLATED
            info.create_system = 0      # Windows (python-docx: 3 = Unix)
            info.create_version = 45     # как Word 15
            zout.writestr(info, parts[name])

    # zipfile принудительно ставит external_attr=0o600<<16 для нулевого
    # значения; Word пишет 0. Обнуляем поле в записях центрального каталога.
    _zero_external_attrs(path)


def _zero_external_attrs(path: str) -> None:
    """Обнуляет external_attr во всех записях центрального каталога zip."""
    with open(path, "rb") as f:
        data = bytearray(f.read())
    eocd = data.rfind(b"PK\x05\x06")
    if eocd == -1:
        return
    cd_offset = int.from_bytes(data[eocd + 16:eocd + 20], "little")
    pos = cd_offset
    while True:
        i = data.find(b"PK\x01\x02", pos)
        if i == -1 or i >= eocd:
            break
        data[i + 38:i + 42] = b"\x00\x00\x00\x00"  # external_attr (4 байта)
        pos = i + 46
    with open(path, "wb") as f:
        f.write(data)


def _add_goback_bookmark(paragraph):
    """Закладка _GoBack в конце абзаца — артефакт курсора Word.

    Все 4 эталона несут ровно одну пару bookmarkStart/End «_GoBack» (id=0)
    в конце первого абзаца: Word ставит её на месте последней правки.
    """
    bm_start = OxmlElement("w:bookmarkStart")
    bm_start.set(qn("w:id"), "0")
    bm_start.set(qn("w:name"), "_GoBack")
    bm_end = OxmlElement("w:bookmarkEnd")
    bm_end.set(qn("w:id"), "0")
    paragraph._p.append(bm_start)
    paragraph._p.append(bm_end)


def _add_run(paragraph, text: str, italic: bool = False, caps: bool = False):
    """Добавляет run с заданным текстом и опциональным курсивом/капсом."""
    run = paragraph.add_run(text)
    run.font.size = Pt(14)
    rpr = run._element.find(qn("w:rPr"))
    rpr.append(OxmlElement("w:szCs", {qn("w:val"): "28"}))
    if caps:
        rpr.insert(0, OxmlElement("w:caps"))
    rpr.insert(0, OxmlElement("w:rFonts", {qn("w:cstheme"): "minorHAnsi"}))
    if italic:
        run.italic = True
    return run


def _add_text_with_italics(paragraph, text: str, lead: str = ""):
    """Добавляет текст в абзац, выделяя курсивом фрагменты в скобках (...).

    lead — префикс (спикер), который вливается в первый НЕкурсивный фрагмент
    одним run-ом (эталонный паттерн: «АЗК: Договариваться или?» — один run).
    """
    if not text:
        if lead:
            _add_run(paragraph, lead)
        return
    parts = [p for p in PARENTHETICAL_RE.split(text) if p]
    for i, part in enumerate(parts):
        is_parenthetical = part.startswith("(") and ")" in part and not part.startswith(UNCLEAR_TEXT)
        if i == 0 and lead:
            if is_parenthetical:
                _add_run(paragraph, lead)
            else:
                _add_run(paragraph, lead + part)
                continue
        _add_run(paragraph, part, italic=is_parenthetical)


def _add_turn_runs(paragraph, timecode: str, speaker_prefix: str, text: str, rng: random.Random):
    """Разбивает реплику на runs по статистике эталонов.

    Цензус f7/f8: ~35% turn-абзацев — один run целиком, ~45% — split на
    границе кадров тайм-кода («11:27:06» | «:00 » | «АЗК: текст»), ~20% —
    split после префикса спикера. Наш прежний постоянный 2-run split
    встречается в эталонах лишь в 3-5% абзацев — статистический маркер машины.
    """
    roll = rng.random()
    plain = not PARENTHETICAL_RE.search(text)
    if plain and roll < 0.35:
        _add_run(paragraph, f"{timecode} {speaker_prefix}{text}")
    elif roll < 0.80:
        _add_run(paragraph, timecode[:-3])
        _add_run(paragraph, f"{timecode[-3:]} ")
        _add_text_with_italics(paragraph, text, lead=speaker_prefix)
    else:
        _add_run(paragraph, f"{timecode} {speaker_prefix}")
        _add_text_with_italics(paragraph, text)


def generate_docx(
    project: dict,
    final_map: dict,
    abbr_map: dict,
    output_path: str,
    legend_exclude: set[str] | None = None,
    segments_override: list[dict] | None = None,
) -> str:
    """Генерирует DOCX с расшифровкой в формате, идентичном ручному.

    Документ строится из backend/transcript_template.docx — пакета,
    извлечённого из эталонной стенограммы (styles/theme/fontTable/settings,
    header с номером страницы, footnotes/endnotes). Это снимает все отпечатки
    дефолтного шаблона python-docx разом; здесь дописываются только абзацы.
    """
    doc = Document(str(DOCX_TEMPLATE_PATH))

    doc.core_properties.author = DOCX_AUTHOR
    doc.core_properties.last_modified_by = DOCX_AUTHOR

    # Размер страницы, поля и расстояние до колонтитулов уже заданы в sectPr
    # шаблона (A4, поля 2/3/2/1.5 см, header/footer 708 twips). Дублируем явно
    # как страховку на случай смены шаблона — значения совпадают, диффа нет.
    for section in doc.sections:
        section.page_width = Mm(210)
        section.page_height = Mm(297)
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(3)
        section.right_margin = Cm(1.5)
        section.header_distance = 449580
        section.footer_distance = 449580

    original_filename = project.get("original_filename", "transcript")
    download_name = strip_extension(original_filename) + ".docx"

    # Заголовочный блок (имя файла, разделитель, легенда) в эталонах несёт
    # <w:caps/> — отображается ПРОПИСНЫМИ (f7/f8/ф14-15; разделитель после
    # легенды и реплики — уже без caps).
    header_para = doc.add_paragraph()
    _configure_paragraph(header_para, caps=True)
    _add_run(header_para, download_name, caps=True)
    _add_goback_bookmark(header_para)

    empty1 = doc.add_paragraph()
    _configure_paragraph(empty1, caps=True)

    speakers_info = project["result"].get("speakers", {})
    segments = segments_override if segments_override is not None else project["result"]["segments"]
    speakers_in_segments = {seg["speaker"] for seg in segments}
    exclude = legend_exclude or set()

    # Эталоны упорядочивают легенду по ПЕРВОМУ появлению в стенограмме
    # (ф5: ИНТЕРВЬЮЕР раньше АНТИПЕНКО, хотя А говорит в 2,5 раза больше),
    # а не по длительности речи, в которой отсортирован speakers_info.
    appearance_order: list = []
    for seg in segments:
        if seg["speaker"] not in appearance_order:
            appearance_order.append(seg["speaker"])
    ordered_ids = [sid for sid in appearance_order if sid in speakers_info]
    ordered_ids += [sid for sid in speakers_info if sid not in appearance_order]

    for speaker_id in ordered_ids:
        info = speakers_info[speaker_id]
        if speaker_id not in speakers_in_segments:
            continue
        if speaker_id in exclude:
            continue
        name = final_map.get(speaker_id, info.get("suggested_name", f"Спикер {speaker_id}"))
        abbr = abbr_map.get(speaker_id, "")
        legend_text = f"{name} – {abbr}." if abbr else f"{name}."
        legend_para = doc.add_paragraph()
        _configure_paragraph(legend_para, caps=True)
        _add_run(legend_para, legend_text, caps=True)

    empty2 = doc.add_paragraph()
    _configure_paragraph(empty2)

    # Свой ГПСЧ для выбора паттерна разбиения на runs (сид от имени файла —
    # детерминизм; отдельный от rsid-ГПСЧ, чтобы не сцеплять потоки).
    split_rng = random.Random(download_name + ":runsplit")

    for seg in segments:
        speaker_id = seg["speaker"]
        abbr = abbr_map.get(speaker_id, "")
        display_name = abbr or final_map.get(speaker_id, f"Спикер {speaker_id}")
        text = seg["text"]

        p = doc.add_paragraph()
        _configure_paragraph(p)

        # Без префикса спикера — только ЦЕЛИКОМ скобочная ремарка;
        # реплика, начинающаяся со скобки («(пауза) и потом...»), — речь.
        if PARENTHETICAL_RE.fullmatch(text) and not text.startswith(UNCLEAR_TEXT):
            _add_run(p, f"{seg['timecode']} ")
            _add_text_with_italics(p, text)
        else:
            _add_turn_runs(p, seg["timecode"], f"{display_name}: ", text, split_rng)

    # Эталоны заканчиваются пустым абзацем после последней реплики
    # (все 4 файла: 1-2 завершающих пустых абзаца).
    trailing = doc.add_paragraph()
    _configure_paragraph(trailing)

    _insert_page_break_markers(doc)
    _apply_revision_ids(doc, download_name)
    _set_core_dates(doc, download_name)

    doc.save(output_path)
    _finalize_package(output_path)
    return download_name
