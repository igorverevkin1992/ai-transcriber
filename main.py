import os
import time
import re
import subprocess
import logging
import uuid
from pathlib import Path
from contextlib import asynccontextmanager
from typing import List

import requests
import boto3
import ffmpeg
from fastapi import FastAPI, HTTPException, BackgroundTasks
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse
from pydantic import BaseModel
from docx import Document
from docx.shared import Pt
from dotenv import load_dotenv

# --- КОНФИГУРАЦИЯ ---
load_dotenv()

logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s [%(levelname)s] %(message)s",
)
logger = logging.getLogger("abtgs")

YANDEX_API_KEY = os.getenv("YANDEX_API_KEY")
AWS_ACCESS_KEY_ID = os.getenv("AWS_ACCESS_KEY_ID")
AWS_SECRET_ACCESS_KEY = os.getenv("AWS_SECRET_ACCESS_KEY")
BUCKET_NAME = os.getenv("BUCKET_NAME", "tv-source-files-2026")
REGION = "ru-central1"

TEMP_DIR = Path("temp_files")
TEMP_DIR.mkdir(exist_ok=True)

# --- S3 КЛИЕНТ ---

def _create_s3_client():
    """Создает S3-клиент. Возвращает None если нет ключей."""
    if not AWS_ACCESS_KEY_ID or not AWS_SECRET_ACCESS_KEY:
        logger.warning("AWS/S3 ключи не найдены — S3 клиент недоступен.")
        return None
    session = boto3.session.Session()
    return session.client(
        service_name="s3",
        endpoint_url="https://storage.yandexcloud.net",
        aws_access_key_id=AWS_ACCESS_KEY_ID,
        aws_secret_access_key=AWS_SECRET_ACCESS_KEY,
        region_name=REGION,
    )

s3 = _create_s3_client()


# --- LIFESPAN ---

@asynccontextmanager
async def lifespan(app: FastAPI):
    """Проверки при старте и очистка при завершении."""
    logger.info("--- ЗАПУСК ПРОВЕРОК ---")

    # Проверка ключей
    if not YANDEX_API_KEY:
        logger.warning("YANDEX_API_KEY не задан — распознавание речи не будет работать.")

    # Проверка S3
    if s3:
        try:
            s3.head_bucket(Bucket=BUCKET_NAME)
            logger.info("S3: бакет '%s' доступен.", BUCKET_NAME)
        except Exception as e:
            logger.error("S3: не удалось получить доступ к бакету '%s': %s", BUCKET_NAME, e)
    else:
        logger.warning("S3 клиент не создан — загрузка файлов в облако невозможна.")

    # Проверка FFmpeg
    try:
        result = subprocess.run(
            ["ffmpeg", "-version"],
            capture_output=True,
            timeout=5,
        )
        if result.returncode == 0:
            logger.info("FFmpeg найден.")
        else:
            logger.error("FFmpeg вернул код ошибки %d.", result.returncode)
    except FileNotFoundError:
        logger.error("FFmpeg не найден в PATH. Установите ffmpeg.")
    except Exception as e:
        logger.error("Ошибка при проверке FFmpeg: %s", e)

    logger.info("--- ПРОВЕРКИ ЗАВЕРШЕНЫ ---")
    yield

    # Очистка temp-файлов при завершении
    for f in TEMP_DIR.iterdir():
        try:
            f.unlink()
        except OSError:
            pass


app = FastAPI(lifespan=lifespan)

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Хранилище проектов (в памяти)
projects_db: dict = {}


# --- ВСПОМОГАТЕЛЬНЫЕ КЛАССЫ ---

class ProjectStatus:
    QUEUED = "В очереди"
    DOWNLOADING = "Скачивание"
    CONVERTING = "Конвертация"
    UPLOADING = "Загрузка в облако"
    TRANSCRIBING = "Распознавание"
    COMPLETED = "Готово"
    ERROR = "Ошибка"


class SpeakerMapping(BaseModel):
    speaker_label: str
    mapped_name: str


class ExportRequest(BaseModel):
    mappings: List[SpeakerMapping]
    filename: str


class CreateProjectRequest(BaseModel):
    url: str


# --- УТИЛИТЫ ---

FILENAME_STOP_WORDS = {
    "лайф", "лайфы", "интер", "синхрон", "снх", "бз",
    "f8", "wav", "mp3", "mp4", "mov", "wmv", "mxf",
}


def parse_filename_metadata(filename: str) -> dict:
    """Извлекает имена спикеров и стартовый таймкод из названия файла."""
    result = {"speakers": [], "start_tc": "00:00:00:00"}

    tc_match = re.search(r"(\d{2}:\d{2}:\d{2}:\d{2})", filename)
    if tc_match:
        result["start_tc"] = tc_match.group(1)
        filename = filename.replace(result["start_tc"], "")

    clean_name = re.sub(r"\.[^.]+$", "", filename)  # убрать расширение
    parts = re.split(r"[,_]+", clean_name)

    for part in parts:
        word = part.strip()
        if (
            word
            and word.lower() not in FILENAME_STOP_WORDS
            and not re.match(r"\d{2}\.\d{2}\.\d{4}", word)
        ):
            result["speakers"].append(word)

    return result


def frames_to_tc(frames: int) -> str:
    """Конвертирует кадры в SMPTE таймкод (25 fps)."""
    h = frames // (25 * 3600)
    rem = frames % (25 * 3600)
    m = rem // (25 * 60)
    rem = rem % (25 * 60)
    s = rem // 25
    f = rem % 25
    return f"{h:02d}:{m:02d}:{s:02d}:{f:02d}"


def tc_to_frames(tc_str: str) -> int:
    """Конвертирует SMPTE таймкод в кадры (25 fps)."""
    try:
        parts = list(map(int, tc_str.split(":")))
        return (parts[0] * 3600 + parts[1] * 60 + parts[2]) * 25 + parts[3]
    except (ValueError, IndexError):
        return 0


def strip_extension(filename: str) -> str:
    """Убирает расширение файла."""
    return re.sub(r"\.[^.]+$", "", filename)


# --- ОСНОВНАЯ ЛОГИКА ОБРАБОТКИ ---

def process_video_task(project_id: str, disk_url: str):
    """Фоновая задача: скачивание -> конвертация -> загрузка в S3 -> распознавание."""
    local_video_path = TEMP_DIR / f"{project_id}_video"
    local_audio_path = TEMP_DIR / f"{project_id}.opus"
    object_name = f"{project_id}.opus"

    try:
        # 1. СКАЧИВАНИЕ
        projects_db[project_id]["status"] = ProjectStatus.DOWNLOADING
        logger.info("[%s] Скачивание файла с Яндекс.Диска...", project_id[:8])

        api_url = "https://cloud-api.yandex.net/v1/disk/public/resources/download"
        resp = requests.get(api_url, params={"public_key": disk_url}, timeout=30)
        resp.raise_for_status()
        download_url = resp.json()["href"]

        # Получаем оригинальное имя файла
        original_filename = "video_source.mp4"
        try:
            meta_url = "https://cloud-api.yandex.net/v1/disk/public/resources"
            meta_resp = requests.get(meta_url, params={"public_key": disk_url}, timeout=15)
            if meta_resp.status_code == 200:
                original_filename = meta_resp.json().get("name", original_filename)
        except requests.RequestException as e:
            logger.warning("[%s] Не удалось получить метаданные файла: %s", project_id[:8], e)

        with requests.get(download_url, stream=True, timeout=600) as r:
            r.raise_for_status()
            with open(local_video_path, "wb") as f:
                for chunk in r.iter_content(chunk_size=65536):
                    f.write(chunk)

        logger.info("[%s] Файл скачан: %s", project_id[:8], original_filename)

        # 2. КОНВЕРТАЦИЯ (FFmpeg)
        projects_db[project_id]["status"] = ProjectStatus.CONVERTING
        logger.info("[%s] Конвертация в OPUS...", project_id[:8])

        (
            ffmpeg
            .input(str(local_video_path))
            .output(str(local_audio_path), acodec="libopus", ac=1, ar=48000)
            .overwrite_output()
            .run(quiet=True)
        )

        logger.info("[%s] Конвертация завершена.", project_id[:8])

        # 3. ЗАГРУЗКА В S3
        if not s3:
            raise RuntimeError("S3 клиент не настроен. Проверьте AWS_ACCESS_KEY_ID и AWS_SECRET_ACCESS_KEY.")

        projects_db[project_id]["status"] = ProjectStatus.UPLOADING
        logger.info("[%s] Загрузка в S3...", project_id[:8])

        s3.upload_file(str(local_audio_path), BUCKET_NAME, object_name)

        file_uri = s3.generate_presigned_url(
            "get_object",
            Params={"Bucket": BUCKET_NAME, "Key": object_name},
            ExpiresIn=3600,
        )

        logger.info("[%s] Файл загружен в S3.", project_id[:8])

        # 4. РАСПОЗНАВАНИЕ (SpeechKit Long Running)
        if not YANDEX_API_KEY:
            raise RuntimeError("YANDEX_API_KEY не задан. Распознавание невозможно.")

        projects_db[project_id]["status"] = ProjectStatus.TRANSCRIBING
        logger.info("[%s] Отправка на распознавание в SpeechKit...", project_id[:8])

        sk_body = {
            "config": {
                "specification": {
                    "languageCode": "ru-RU",
                    "literature_text": True,
                    "profanityFilter": False,
                    "audioEncoding": "OGG_OPUS",
                    "sampleRateHertz": 48000,
                    "audioChannelCount": 1,
                }
            },
            "audio": {
                "uri": file_uri,
            },
        }

        sk_headers = {"Authorization": f"Api-Key {YANDEX_API_KEY}"}
        sk_resp = requests.post(
            "https://transcribe.api.cloud.yandex.net/speech/stt/v2/longRunningRecognize",
            json=sk_body,
            headers=sk_headers,
            timeout=60,
        )
        sk_resp.raise_for_status()
        operation_id = sk_resp.json()["id"]
        logger.info("[%s] Операция SpeechKit: %s", project_id[:8], operation_id)

        # Ожидание результата
        poll_interval = 3
        max_polls = 600  # ~30 минут максимум
        for _ in range(max_polls):
            time.sleep(poll_interval)
            op_resp = requests.get(
                f"https://operation.api.cloud.yandex.net/operations/{operation_id}",
                headers=sk_headers,
                timeout=30,
            )
            op_data = op_resp.json()
            if op_data.get("done"):
                break
        else:
            raise TimeoutError("SpeechKit: превышено время ожидания распознавания.")

        if "error" in op_data:
            raise RuntimeError(f"SpeechKit Error: {op_data['error']}")

        # 5. ОБРАБОТКА РЕЗУЛЬТАТА
        chunks = op_data["response"]["chunks"]
        meta = parse_filename_metadata(original_filename)
        projects_db[project_id]["original_filename"] = original_filename

        speaker_durations: dict[str, float] = {}
        raw_segments = []
        start_frames = tc_to_frames(meta["start_tc"])

        for chunk in chunks:
            channel = str(chunk.get("channelTag", "1"))
            alternatives = chunk.get("alternatives", [])
            if not alternatives:
                continue

            alt = alternatives[0]
            text = alt.get("text", "")
            words = alt.get("words", [])
            if not words:
                continue

            start_s_str = words[0].get("startTime", "0s")
            end_s_str = words[-1].get("endTime", "0s")

            start_s = float(start_s_str.replace("s", ""))
            end_s = float(end_s_str.replace("s", ""))

            dur = end_s - start_s
            speaker_durations[channel] = speaker_durations.get(channel, 0) + dur

            abs_frames = start_frames + int(start_s * 25)
            tc_formatted = frames_to_tc(abs_frames)

            raw_segments.append({
                "timecode": tc_formatted,
                "speaker": channel,
                "text": text,
            })

        # Назначение имен спикерам
        detected_speakers = {}
        sorted_voices = sorted(speaker_durations.items(), key=lambda x: x[1], reverse=True)
        file_names = meta["speakers"]

        for i, (voice_id, dur) in enumerate(sorted_voices):
            suggested = f"Спикер {voice_id}"
            if i < len(file_names):
                suggested = file_names[i]

            detected_speakers[voice_id] = {
                "duration_sec": round(dur, 1),
                "suggested_name": suggested,
            }

        projects_db[project_id]["result"] = {
            "segments": raw_segments,
            "speakers": detected_speakers,
            "meta": {**meta, "original_filename": original_filename},
        }
        projects_db[project_id]["status"] = ProjectStatus.COMPLETED
        logger.info("[%s] Обработка завершена. Сегментов: %d, Спикеров: %d",
                     project_id[:8], len(raw_segments), len(detected_speakers))

    except Exception as e:
        logger.exception("[%s] Ошибка обработки: %s", project_id[:8], e)
        projects_db[project_id]["status"] = ProjectStatus.ERROR
        projects_db[project_id]["error"] = str(e)

    finally:
        # Очистка временных файлов
        for path in (local_video_path, local_audio_path):
            try:
                if path.exists():
                    path.unlink()
            except OSError:
                pass
        # Удаление из S3
        if s3:
            try:
                s3.delete_object(Bucket=BUCKET_NAME, Key=object_name)
            except Exception as e:
                logger.warning("[%s] Не удалось удалить %s из S3: %s", project_id[:8], object_name, e)


# --- API ENDPOINTS ---

@app.get("/")
def read_root():
    """Проверка здоровья сервера."""
    return {
        "status": "ok",
        "service": "ABTGS Backend",
        "message": "Перейдите на http://localhost:3000 для работы с интерфейсом.",
    }


@app.post("/api/v1/projects")
async def create_project(req: CreateProjectRequest, background_tasks: BackgroundTasks):
    """Создает проект и запускает фоновую обработку."""
    pid = str(uuid.uuid4())
    projects_db[pid] = {
        "id": pid,
        "status": ProjectStatus.QUEUED,
        "created_at": time.time(),
    }
    background_tasks.add_task(process_video_task, pid, req.url)
    logger.info("Проект создан: %s для URL: %s", pid[:8], req.url[:60])
    return {"id": pid}


@app.get("/api/v1/projects/{pid}/status")
async def get_status(pid: str):
    """Возвращает текущий статус обработки проекта."""
    if pid not in projects_db:
        raise HTTPException(status_code=404, detail="Проект не найден")
    proj = projects_db[pid]
    return {
        "status": proj["status"],
        "error": proj.get("error"),
    }


@app.get("/api/v1/projects/{pid}")
async def get_result(pid: str):
    """Возвращает результаты распознавания."""
    proj = projects_db.get(pid)
    if not proj:
        raise HTTPException(status_code=404, detail="Проект не найден")
    if proj["status"] != ProjectStatus.COMPLETED:
        raise HTTPException(status_code=400, detail="Обработка ещё не завершена")
    return proj["result"]


@app.post("/api/v1/projects/{pid}/export")
async def export_docx(pid: str, req: ExportRequest, background_tasks: BackgroundTasks):
    """Генерирует DOCX-файл с транскриптом и отдает для скачивания."""
    proj = projects_db.get(pid)
    if not proj:
        raise HTTPException(status_code=404, detail="Проект не найден")

    if "result" not in proj:
        raise HTTPException(status_code=400, detail="Результаты обработки отсутствуют")

    final_map = {m.speaker_label: m.mapped_name for m in req.mappings}

    doc = Document()
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Times New Roman"
    font.size = Pt(12)

    original_filename = proj.get("original_filename", "transcript")

    # Заголовок документа
    header_para = doc.add_paragraph()
    header_run = header_para.add_run(f"ИСХОДНИК: {original_filename}")
    header_run.bold = True
    doc.add_paragraph()  # Пустая строка после заголовка

    # Список спикеров
    speakers_info = proj["result"].get("speakers", {})
    for speaker_id, info in speakers_info.items():
        name = final_map.get(speaker_id, info.get("suggested_name", f"Спикер {speaker_id}"))
        doc.add_paragraph(f"{name.upper()}")
    doc.add_paragraph()  # Пустая строка после списка спикеров

    # Сегменты
    segments = proj["result"]["segments"]
    for seg in segments:
        speaker_name = final_map.get(seg["speaker"], f"Спикер {seg['speaker']}")
        p = doc.add_paragraph()
        p_runner = p.add_run(f"{seg['timecode']} {speaker_name}: ")
        p_runner.bold = True
        p.add_run(seg["text"])

    output_path = TEMP_DIR / f"transcript_{pid}.docx"
    doc.save(str(output_path))

    # Формируем имя файла для скачивания
    download_name = strip_extension(original_filename) + ".docx"

    # Удалить файл после отправки
    background_tasks.add_task(os.unlink, str(output_path))

    return FileResponse(
        path=str(output_path),
        filename=download_name,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )


if __name__ == "__main__":
    import uvicorn

    logger.info("Запуск ABTGS Backend на http://localhost:8000")
    uvicorn.run(app, host="0.0.0.0", port=8000)
